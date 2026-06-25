"""
=============================================================================
GENERATOR LAPORAN ILMIAH LENGKAP - PROJECT DATA WAREHOUSE
=============================================================================
Nama   : Nathanael Komang Bagus Prakarsa
NIM    : 245091107111005
Prodi  : Sains Data - Departemen Statistika
Fakultas: Sains, Teknologi, dan Matematika
Univ   : Universitas Brawijaya
=============================================================================
"""

import os
import sys
import io
import sqlite3
import urllib.request
import urllib.error

# Auto-install dependencies
def install_if_missing(pkg, import_name=None):
    import_name = import_name or pkg
    try:
        __import__(import_name)
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

install_if_missing("python-docx", "docx")
install_if_missing("matplotlib")
install_if_missing("seaborn")
install_if_missing("pandas")
install_if_missing("numpy")
install_if_missing("Pillow", "PIL")
install_if_missing("requests")

import requests
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import seaborn as sns
from PIL import Image

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# =============================================================================
# KONFIGURASI PATH
# =============================================================================
BASE_DIR = r"d:\Nael\Mata Kuliah Sains Data\.vscode\Kuliah\Semester 4\DWH Dosen"
DB_PATH  = os.path.join(BASE_DIR, "dwh_nilai.db")
OUTPUT_PATH = os.path.join(BASE_DIR, "Project DW_Nathanael Komang Bagus P_245091107111005.docx")

STAR_SCHEMA_IMG   = os.path.join(BASE_DIR, "Star Schema.png")
ALUR_DWH_IMG      = os.path.join(BASE_DIR, "DWH Alur drawio.png")
DASHBOARD_IMG     = os.path.join(BASE_DIR, "_dashboard_combined.png")
UB_LOGO_PATH      = os.path.join(BASE_DIR, "_ub_logo.png")

# =============================================================================
# WARNA DESAIN (Dark Mode Premium)
# =============================================================================
BG_COLOR      = '#0F0F13'
CARD_COLOR    = '#1A1A24'
PANEL_COLOR   = '#22222E'
TEXT_COLOR    = '#E8E8F0'
ACCENT_BLUE   = '#4F8EF7'
ACCENT_PURPLE = '#9B72F5'
ACCENT_GREEN  = '#22D3A3'
ACCENT_YELLOW = '#FBBF24'
ACCENT_RED    = '#F87171'
ACCENT_CYAN   = '#38BDF8'
ACCENT_ORANGE = '#FB923C'
GRID_COLOR    = '#2A2A38'

# =============================================================================
# STEP 1: UNDUH LOGO UNIVERSITAS BRAWIJAYA
# =============================================================================
def download_ub_logo():
    """Mengunduh logo Universitas Brawijaya dari web resmi."""
    logo_url = "https://upload.wikimedia.org/wikipedia/id/2/2d/Lambang_Universitas_Brawijaya.png"
    if os.path.exists(UB_LOGO_PATH):
        print(f"Logo UB sudah ada: {UB_LOGO_PATH}")
        return True
    print("Mengunduh logo Universitas Brawijaya...")
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        req = urllib.request.Request(logo_url, headers=headers)
        with urllib.request.urlopen(req, timeout=15) as response:
            data = response.read()
        with open(UB_LOGO_PATH, 'wb') as f:
            f.write(data)
        print(f"Logo UB berhasil diunduh: {UB_LOGO_PATH}")
        return True
    except Exception as e:
        print(f"Gagal mengunduh logo: {e}. Akan dibuat logo placeholder.")
        _create_ub_logo_placeholder()
        return True

def _create_ub_logo_placeholder():
    """Membuat placeholder logo UB jika unduhan gagal."""
    fig, ax = plt.subplots(figsize=(2, 2), facecolor='white')
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.axis('off')
    circle = plt.Circle((0.5, 0.5), 0.45, color='#003366', fill=True)
    ax.add_patch(circle)
    ax.text(0.5, 0.55, 'UB', ha='center', va='center', fontsize=28,
            fontweight='bold', color='white')
    ax.text(0.5, 0.35, 'BRAWIJAYA', ha='center', va='center', fontsize=7,
            color='#FFD700', fontweight='bold')
    plt.savefig(UB_LOGO_PATH, dpi=150, bbox_inches='tight',
                facecolor='white')
    plt.close()


# =============================================================================
# STEP 2: BUAT DASHBOARD CHART TERPADU (5 chart dalam 1 figure)
# =============================================================================
def build_dashboard_chart():
    """Membuat satu figure dashboard yang menggabungkan semua 5 visualisasi."""
    print("Membangun dashboard chart terpadu...")

    conn = sqlite3.connect(DB_PATH)

    # ---- Ambil semua data yang dibutuhkan ----
    # Q1: Progress kelengkapan
    df1 = pd.read_sql_query("""
        SELECT tipe_kelas, status_nilai, COUNT(*) AS jumlah
        FROM Fact_Nilai
        GROUP BY tipe_kelas, status_nilai
    """, conn)

    # Q2: Teori vs Praktikum distribusi nilai
    df2 = pd.read_sql_query("""
        SELECT tipe_kelas, nilai_angka
        FROM Fact_Nilai WHERE status_nilai = 'lengkap'
    """, conn)

    # Q3: Rata-rata per matkul & tipe
    df3 = pd.read_sql_query("""
        SELECT m.nama_matkul, f.tipe_kelas,
               ROUND(AVG(f.nilai_angka), 2) AS rata_rata
        FROM Fact_Nilai f
        JOIN Dim_MataKuliah m ON f.matkul_id = m.matkul_id
        WHERE f.status_nilai = 'lengkap'
        GROUP BY m.nama_matkul, f.tipe_kelas
        ORDER BY rata_rata DESC
    """, conn)

    # Q4: Distribusi komponen praktikum
    df4 = pd.read_sql_query("""
        SELECT m.nama_matkul, k.nama_komponen, f.nilai_angka
        FROM Fact_Nilai f
        JOIN Dim_MataKuliah m ON f.matkul_id = m.matkul_id
        JOIN Dim_KomponenNilai k ON f.komponen_id = k.komponen_id
        WHERE f.status_nilai = 'lengkap' AND f.tipe_kelas = 'praktikum'
    """, conn)

    # Q5: Top 10 nilai akhir praktikum
    df5 = pd.read_sql_query("""
        SELECT m.nama AS nama_mahasiswa, mk.nama_matkul,
               ROUND(SUM(f.nilai_angka * (k.bobot_persen / 100.0)), 2) AS nilai_akhir
        FROM Fact_Nilai f
        JOIN Dim_Mahasiswa m ON f.mahasiswa_id = m.mahasiswa_id
        JOIN Dim_MataKuliah mk ON f.matkul_id = mk.matkul_id
        JOIN Dim_KomponenNilai k ON f.komponen_id = k.komponen_id
        WHERE f.tipe_kelas = 'praktikum' AND f.nilai_angka IS NOT NULL
        GROUP BY m.nim, m.nama, mk.nama_matkul
        ORDER BY nilai_akhir DESC LIMIT 10
    """, conn)

    # Q: Summary metrics
    df_metrics = pd.read_sql_query("""
        SELECT
            (SELECT COUNT(DISTINCT mahasiswa_id) FROM Dim_Mahasiswa) AS total_mhs,
            (SELECT COUNT(*) FROM Fact_Nilai) AS total_records,
            (SELECT COUNT(*) FROM Fact_Nilai WHERE status_nilai='lengkap') AS total_lengkap,
            (SELECT ROUND(AVG(nilai_angka),2) FROM Fact_Nilai WHERE status_nilai='lengkap') AS avg_nilai,
            (SELECT ROUND(AVG(nilai_angka),2) FROM Fact_Nilai
             WHERE status_nilai='lengkap' AND tipe_kelas='praktikum') AS avg_prak,
            (SELECT ROUND(AVG(nilai_angka),2) FROM Fact_Nilai
             WHERE status_nilai='lengkap' AND tipe_kelas='teori') AS avg_teori
    """, conn)
    conn.close()

    metrics = df_metrics.iloc[0]

    # ---- Setup Figure ----
    plt.rcParams.update({
        'figure.facecolor': BG_COLOR,
        'axes.facecolor': CARD_COLOR,
        'text.color': TEXT_COLOR,
        'axes.labelcolor': TEXT_COLOR,
        'xtick.color': TEXT_COLOR,
        'ytick.color': TEXT_COLOR,
        'font.family': 'DejaVu Sans',
        'grid.color': GRID_COLOR,
        'grid.linestyle': '--',
        'grid.linewidth': 0.5,
        'axes.edgecolor': GRID_COLOR,
        'axes.spines.top': False,
        'axes.spines.right': False,
    })

    fig = plt.figure(figsize=(22, 16), facecolor=BG_COLOR)

    # --- HEADER DASHBOARD ---
    header_ax = fig.add_axes([0, 0.935, 1, 0.065])
    header_ax.set_facecolor('#1C1C2E')
    header_ax.axis('off')
    header_ax.add_patch(FancyBboxPatch((0, 0), 1, 1, transform=header_ax.transAxes,
                         boxstyle="square,pad=0", facecolor='#1C1C2E', zorder=0))
    # Gradient-like side accent
    for i, (x_start, clr) in enumerate([(0, ACCENT_BLUE), (0.33, ACCENT_PURPLE), (0.66, ACCENT_GREEN)]):
        header_ax.axvline(x=x_start, color=clr, linewidth=2, alpha=0.3)
    header_ax.text(0.5, 0.65, '📊  DASHBOARD ANALITIK DATA WAREHOUSE NILAI AKADEMIK MAHASISWA',
                   ha='center', va='center', fontsize=16, fontweight='bold',
                   color='white', transform=header_ax.transAxes)
    header_ax.text(0.5, 0.2,
                   'Program Studi Sains Data  |  Departemen Statistika  |  Universitas Brawijaya  |  T.A. 2025/2026',
                   ha='center', va='center', fontsize=9, color='#9999BB',
                   transform=header_ax.transAxes)

    # --- METRIC CARDS ROW ---
    card_metrics = [
        ('👥 Total Mahasiswa', f"{int(metrics['total_mhs'])}", ACCENT_BLUE),
        ('📋 Total Records', f"{int(metrics['total_records'])}", ACCENT_PURPLE),
        ('✅ Nilai Lengkap', f"{int(metrics['total_lengkap'])}", ACCENT_GREEN),
        ('📈 Rata-rata Praktikum', f"{metrics['avg_prak']:.2f}", ACCENT_YELLOW),
        ('📉 Rata-rata Teori', f"{metrics['avg_teori']:.2f}", ACCENT_CYAN),
        ('⭐ Rata-rata Overall', f"{metrics['avg_nilai']:.2f}", ACCENT_ORANGE),
    ]

    card_width  = 1 / 6
    card_height = 0.08
    card_y      = 0.855

    for i, (label, value, color) in enumerate(card_metrics):
        x = i * card_width
        card_ax = fig.add_axes([x + 0.003, card_y, card_width - 0.006, card_height])
        card_ax.set_facecolor(PANEL_COLOR)
        card_ax.axis('off')
        # Top color bar
        card_ax.add_patch(FancyBboxPatch((0, 0.85), 1, 0.15, transform=card_ax.transAxes,
                           boxstyle="square,pad=0", facecolor=color, alpha=0.8))
        card_ax.text(0.5, 0.52, value, ha='center', va='center', fontsize=20,
                     fontweight='bold', color=color, transform=card_ax.transAxes)
        card_ax.text(0.5, 0.18, label, ha='center', va='center', fontsize=8,
                     color='#AAAACC', transform=card_ax.transAxes)

    # ---- GRIDSPEC for 5 charts ----
    gs = gridspec.GridSpec(2, 3,
                           left=0.04, right=0.97,
                           top=0.84, bottom=0.04,
                           hspace=0.38, wspace=0.32)

    # ==== CHART 1: Progress Kelengkapan (Stacked Horizontal Bar) ====
    ax1 = fig.add_subplot(gs[0, 0])
    ax1.set_facecolor(CARD_COLOR)

    df1_pivot = df1.pivot(index='tipe_kelas', columns='status_nilai', values='jumlah').fillna(0)
    for col in ['lengkap', 'belum_dikoreksi', 'belum_dilaksanakan']:
        if col not in df1_pivot.columns:
            df1_pivot[col] = 0
    df1_pivot = df1_pivot[['lengkap', 'belum_dikoreksi', 'belum_dilaksanakan']]
    df1_pct = df1_pivot.div(df1_pivot.sum(axis=1), axis=0) * 100

    colors_chart1 = [ACCENT_GREEN, ACCENT_YELLOW, ACCENT_RED]
    lefts = np.zeros(len(df1_pct))
    bar_h = 0.45
    y_pos = np.arange(len(df1_pct))

    for j, (col, color) in enumerate(zip(df1_pct.columns, colors_chart1)):
        vals = df1_pct[col].values
        bars = ax1.barh(y_pos, vals, left=lefts, color=color, height=bar_h,
                        edgecolor=BG_COLOR, linewidth=0.8, label=col.replace('_', ' ').title())
        for bar_idx, (v, l) in enumerate(zip(vals, lefts)):
            if v > 8:
                ax1.text(l + v/2, bar_idx, f"{v:.1f}%",
                         ha='center', va='center', color='white',
                         fontsize=9, fontweight='bold')
        lefts += vals

    ax1.set_yticks(y_pos)
    ax1.set_yticklabels(['Praktikum', 'Teori'], fontsize=10, fontweight='bold')
    ax1.set_xlim(0, 100)
    ax1.set_xlabel("Persentase (%)", fontsize=9)
    ax1.set_title("Progress Kelengkapan Data Nilai", fontsize=11, fontweight='bold',
                  pad=12, color=TEXT_COLOR)
    ax1.grid(True, axis='x', alpha=0.4)

    legend1 = ax1.legend(loc='lower right', fontsize=7.5, frameon=True,
                          facecolor=PANEL_COLOR, edgecolor=GRID_COLOR,
                          framealpha=0.9)
    for text in legend1.get_texts():
        text.set_color(TEXT_COLOR)

    # Color accent line top
    ax1.axhline(y=len(y_pos) - 0.5 + bar_h/2 + 0.08, color=ACCENT_BLUE, linewidth=2, alpha=0.5)

    # ==== CHART 2: Violin Plot Teori vs Praktikum ====
    ax2 = fig.add_subplot(gs[0, 1])
    ax2.set_facecolor(CARD_COLOR)

    palette2 = {'praktikum': ACCENT_BLUE, 'teori': ACCENT_PURPLE}
    vp = sns.violinplot(x='tipe_kelas', y='nilai_angka', data=df2,
                        palette=palette2, ax=ax2, inner='quartile',
                        linewidth=1.2, cut=0.5)

    means2 = df2.groupby('tipe_kelas')['nilai_angka'].mean()
    medians2 = df2.groupby('tipe_kelas')['nilai_angka'].median()
    stds2 = df2.groupby('tipe_kelas')['nilai_angka'].std()

    x_order = ['praktikum', 'teori']
    for i, tipe in enumerate(x_order):
        if tipe in means2.index:
            ax2.scatter(i, means2[tipe], color=ACCENT_YELLOW, zorder=5, s=60,
                        marker='D', label='Mean' if i == 0 else '')
            ax2.annotate(
                f"μ={means2[tipe]:.1f}\nσ={stds2[tipe]:.1f}",
                xy=(i, means2[tipe]),
                xytext=(i + 0.28, means2[tipe] + 5),
                fontsize=8.5, color=ACCENT_YELLOW, fontweight='bold',
                arrowprops=dict(arrowstyle='->', color=ACCENT_YELLOW, lw=1.2),
                bbox=dict(boxstyle='round,pad=0.3', facecolor=PANEL_COLOR,
                          edgecolor=ACCENT_YELLOW, alpha=0.85)
            )

    ax2.set_title("Distribusi Nilai: Teori vs Praktikum", fontsize=11, fontweight='bold',
                  pad=12, color=TEXT_COLOR)
    ax2.set_xlabel("Tipe Kelas", fontsize=9)
    ax2.set_ylabel("Nilai Angka", fontsize=9)
    ax2.set_xticklabels(['Praktikum', 'Teori'], fontsize=10, fontweight='bold')
    ax2.grid(True, axis='y', alpha=0.4)
    ax2.set_ylim(0, 115)

    # ==== CHART 3: Grouped Bar - Rata-rata per Matkul ====
    ax3 = fig.add_subplot(gs[0, 2])
    ax3.set_facecolor(CARD_COLOR)

    palette3 = {'praktikum': ACCENT_BLUE, 'teori': ACCENT_PURPLE}

    # Shorten matkul names
    df3['matkul_short'] = df3['nama_matkul'].str.replace('Metode ', 'M.', regex=False)
    df3['matkul_short'] = df3['matkul_short'].str.replace('Statistika', 'Stat.', regex=False)
    df3['matkul_short'] = df3['matkul_short'].str.replace('Sains Data', 'S.D.', regex=False)

    sns.barplot(x='matkul_short', y='rata_rata', hue='tipe_kelas',
                data=df3, palette=palette3, ax=ax3,
                edgecolor=BG_COLOR, linewidth=0.8, width=0.6)

    for p in ax3.patches:
        h = p.get_height()
        if h > 0:
            ax3.annotate(f"{h:.1f}",
                         (p.get_x() + p.get_width()/2., h + 1),
                         ha='center', va='bottom', color=TEXT_COLOR,
                         fontsize=8.5, fontweight='bold')

    ax3.set_ylim(0, 115)
    ax3.set_title("Rata-rata Nilai per Mata Kuliah", fontsize=11, fontweight='bold',
                  pad=12, color=TEXT_COLOR)
    ax3.set_xlabel("Mata Kuliah", fontsize=9)
    ax3.set_ylabel("Rata-rata Nilai", fontsize=9)
    ax3.tick_params(axis='x', labelsize=8)
    ax3.grid(True, axis='y', alpha=0.4)

    legend3 = ax3.legend(title='Tipe Kelas', facecolor=PANEL_COLOR,
                          edgecolor=GRID_COLOR, fontsize=8, title_fontsize=8.5)
    legend3.get_title().set_color(TEXT_COLOR)
    for text in legend3.get_texts():
        text.set_color(TEXT_COLOR)

    # Annotate reference line
    ax3.axhline(y=80, color=ACCENT_YELLOW, linestyle='--', linewidth=1, alpha=0.6)
    ax3.text(ax3.get_xlim()[1] * 0.98, 82, 'KKM=80', ha='right', fontsize=8,
             color=ACCENT_YELLOW, alpha=0.8)

    # ==== CHART 4: Box Plot Komponen Praktikum ====
    ax4 = fig.add_subplot(gs[1, 0:2])
    ax4.set_facecolor(CARD_COLOR)

    # Filter hanya komponen yang ada
    df4_filtered = df4[df4['nilai_angka'].notna()]
    df4_filtered = df4_filtered[df4_filtered['nama_komponen'].isin(
        ['Tugas 1', 'Tugas 2', 'Presensi', 'Sikap', 'UTP', 'UAP']
    )]

    matkul_names = df4_filtered['nama_matkul'].unique()
    palette4 = dict(zip(matkul_names,
                        [ACCENT_BLUE, ACCENT_PURPLE, ACCENT_GREEN][:len(matkul_names)]))

    # Shorten nama_matkul for legend
    df4_filtered = df4_filtered.copy()
    df4_filtered['matkul_abbr'] = df4_filtered['nama_matkul'].apply(
        lambda x: ''.join([w[0] for w in x.split()])
    )
    palette4b = {'MSI': ACCENT_BLUE, 'MSII': ACCENT_PURPLE, 'MSDI': ACCENT_GREEN}
    valid_komps = ['Tugas 1', 'Tugas 2', 'Presensi', 'Sikap', 'UTP', 'UAP']

    sns.boxplot(
        x='nama_komponen', y='nilai_angka',
        hue='nama_matkul',
        data=df4_filtered[df4_filtered['nama_komponen'].isin(valid_komps)],
        order=valid_komps,
        palette=[ACCENT_BLUE, ACCENT_PURPLE, ACCENT_GREEN],
        ax=ax4, linewidth=1.1,
        flierprops=dict(marker='o', markerfacecolor=ACCENT_RED,
                        markersize=4, markeredgecolor=BG_COLOR)
    )

    ax4.set_title("Distribusi Nilai per Komponen Kelas Praktikum", fontsize=11,
                  fontweight='bold', pad=12, color=TEXT_COLOR)
    ax4.set_xlabel("Komponen Nilai", fontsize=9)
    ax4.set_ylabel("Nilai Mahasiswa", fontsize=9)
    ax4.set_ylim(40, 115)
    ax4.grid(True, axis='y', alpha=0.4)

    legend4 = ax4.legend(title='Mata Kuliah', loc='lower right',
                          facecolor=PANEL_COLOR, edgecolor=GRID_COLOR,
                          fontsize=8, title_fontsize=8.5)
    legend4.get_title().set_color(TEXT_COLOR)
    for text in legend4.get_texts():
        text.set_color(TEXT_COLOR)

    ax4.axhline(y=80, color=ACCENT_YELLOW, linestyle='--', linewidth=1, alpha=0.6)

    # ==== CHART 5: Horizontal Bar - Top 10 Mahasiswa ====
    ax5 = fig.add_subplot(gs[1, 2])
    ax5.set_facecolor(CARD_COLOR)

    df5_sorted = df5.sort_values('nilai_akhir', ascending=True).tail(10)
    df5_sorted['label'] = (df5_sorted['nama_mahasiswa'].str.split().str[0] + "\n(" +
                           df5_sorted['nama_matkul'].apply(
                               lambda x: ''.join([w[0] for w in x.split()])) + ")")

    # Color gradient based on value
    norm_vals = (df5_sorted['nilai_akhir'] - df5_sorted['nilai_akhir'].min())
    norm_vals = norm_vals / (norm_vals.max() + 0.01)
    bar_colors = plt.cm.plasma(0.4 + norm_vals.values * 0.5)

    bars5 = ax5.barh(range(len(df5_sorted)), df5_sorted['nilai_akhir'],
                     color=bar_colors, edgecolor=BG_COLOR, linewidth=0.8, height=0.65)

    ax5.set_yticks(range(len(df5_sorted)))
    ax5.set_yticklabels(df5_sorted['label'].values, fontsize=7.5)
    ax5.set_xlim(85, 102)
    ax5.set_title("Top 10 Nilai Akhir Praktikum", fontsize=11, fontweight='bold',
                  pad=12, color=TEXT_COLOR)
    ax5.set_xlabel("Nilai Akhir Estimasi", fontsize=9)
    ax5.grid(True, axis='x', alpha=0.4)

    for bar, val in zip(bars5, df5_sorted['nilai_akhir']):
        ax5.text(val + 0.1, bar.get_y() + bar.get_height()/2,
                 f" {val:.2f}", va='center', ha='left', fontsize=8.5,
                 color=ACCENT_YELLOW, fontweight='bold')

    # --- FOOTER ---
    footer_ax = fig.add_axes([0, 0, 1, 0.018])
    footer_ax.set_facecolor('#12121C')
    footer_ax.axis('off')
    footer_ax.text(0.5, 0.5,
                   'Nathanael Komang Bagus Prakarsa  ·  NIM 245091107111005  ·  '
                   'Program Studi Sains Data  ·  Universitas Brawijaya  ·  © 2026',
                   ha='center', va='center', fontsize=8, color='#666688',
                   transform=footer_ax.transAxes)

    plt.savefig(DASHBOARD_IMG, dpi=200, bbox_inches='tight', facecolor=BG_COLOR)
    plt.close()
    print(f"Dashboard chart berhasil dibuat: {DASHBOARD_IMG}")


# =============================================================================
# HELPER FUNCTIONS UNTUK DOCX
# =============================================================================
def set_cell_bg(cell, hex_color):
    """Menetapkan warna background sel tabel DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Menetapkan border sel tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading_styled(doc, text, level=1, color_hex='1F3864'):
    """Menambahkan heading dengan style kustom."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(color_hex)
        # Tambah garis bawah
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single')
        bottom_bdr.set(qn('w:sz'), '8')
        bottom_bdr.set(qn('w:color'), color_hex)
        pBdr.append(bottom_bdr)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(color_hex)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(color_hex)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_body_text(doc, text, justify=True, indent=False):
    """Menambahkan paragraf teks dengan format standar."""
    para = doc.add_paragraph(text)
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = Pt(18)
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    return para

def add_bullet(doc, text, level=0):
    """Menambahkan bullet/list item."""
    para = doc.add_paragraph(style='List Bullet')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = Pt(18)
    return para

def add_numbered(doc, text, num):
    """Menambahkan numbered list item."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_num = para.add_run(f"{num}. ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.name = 'Times New Roman'
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = Pt(18)
    return para

def set_document_margins(doc):
    """Mengatur margin dokumen sesuai standar akademik Indonesia."""
    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin   = Cm(4)
        section.right_margin  = Cm(3)

def add_page_break(doc):
    doc.add_page_break()

def add_caption(doc, text, fig_num=None):
    """Menambahkan keterangan gambar/tabel."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(70, 70, 70)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(12)
    return para

def add_image_centered(doc, img_path, width_inches=6.0, caption=None):
    """Menambahkan gambar terpusat dalam dokumen."""
    if os.path.exists(img_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        if caption:
            add_caption(doc, caption)
    else:
        add_body_text(doc, f"[Gambar tidak ditemukan: {os.path.basename(img_path)}]")


# =============================================================================
# STEP 3: BUILD DOCX REPORT
# =============================================================================
def build_docx_report():
    """Membangun laporan DOCX lengkap dan komprehensif."""
    print("Membangun laporan DOCX...")
    doc = Document()
    set_document_margins(doc)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # =========================================================================
    # HALAMAN JUDUL
    # =========================================================================
    # Logo UB
    para_logo = doc.add_paragraph()
    para_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = para_logo.add_run()
    if os.path.exists(UB_LOGO_PATH):
        run_logo.add_picture(UB_LOGO_PATH, width=Inches(1.5))
    para_logo.paragraph_format.space_after = Pt(8)

    # Nama universitas
    def add_cover_text(doc, text, size=12, bold=False, color='000000', space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor.from_string(color)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        return p

    add_cover_text(doc, "UNIVERSITAS BRAWIJAYA", size=14, bold=True, color='003366', space_after=2)
    add_cover_text(doc, "FAKULTAS SAINS, TEKNOLOGI, DAN MATEMATIKA", size=11, bold=True, color='003366', space_after=2)
    add_cover_text(doc, "DEPARTEMEN STATISTIKA", size=11, bold=True, color='003366', space_after=2)
    add_cover_text(doc, "PROGRAM STUDI SAINS DATA", size=11, bold=True, color='003366', space_after=10)

    # Garis pemisah
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement('w:bottom')
    bdr.set(qn('w:val'), 'double')
    bdr.set(qn('w:sz'), '6')
    bdr.set(qn('w:color'), '003366')
    pBdr.append(bdr)
    pPr.append(pBdr)
    p_line.paragraph_format.space_after = Pt(12)

    # Judul utama
    add_cover_text(doc, "LAPORAN ILMIAH", size=16, bold=True, color='1F3864', space_after=6)
    add_cover_text(doc, "PROJECT DATA WAREHOUSE", size=18, bold=True, color='003366', space_after=6)
    add_cover_text(doc,
        "Implementasi Data Warehouse untuk Pemantauan Nilai Akademik Mahasiswa\n"
        "Program Studi Sains Data Universitas Brawijaya\n"
        "Tahun Akademik 2025/2026",
        size=12, bold=False, color='444444', space_after=20)

    # Garis tengah
    p_mid = doc.add_paragraph()
    pPr2 = p_mid._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    bdr2 = OxmlElement('w:bottom')
    bdr2.set(qn('w:val'), 'single')
    bdr2.set(qn('w:sz'), '4')
    bdr2.set(qn('w:color'), 'AAAAAA')
    pBdr2.append(bdr2)
    pPr2.append(pBdr2)
    p_mid.paragraph_format.space_after = Pt(16)

    add_cover_text(doc, "Disusun Oleh:", size=11, bold=False, color='555555', space_after=4)
    add_cover_text(doc, "Nathanael Komang Bagus Prakarsa", size=13, bold=True, color='003366', space_after=4)
    add_cover_text(doc, "NIM: 245091107111005", size=12, bold=False, color='333333', space_after=20)

    add_cover_text(doc,
        "Malang\n2026",
        size=12, bold=False, color='333333', space_after=4)

    add_page_break(doc)

    # =========================================================================
    # KATA PENGANTAR
    # =========================================================================
    add_heading_styled(doc, "KATA PENGANTAR", level=1, color_hex='1F3864')
    doc.add_paragraph()

    pengantar_paras = [
        ("Puji dan syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa atas "
         "segala rahmat dan karunia-Nya sehingga laporan ilmiah ini dapat diselesaikan "
         "dengan baik. Laporan berjudul \"Implementasi Data Warehouse untuk Pemantauan "
         "Nilai Akademik Mahasiswa\" ini disusun sebagai pemenuhan tugas Project-Based "
         "Learning Data Warehouse Ke-II pada Program Studi Sains Data, Departemen "
         "Statistika, Fakultas Sains, Teknologi, dan Matematika, Universitas Brawijaya."),

        ("Laporan ini membahas seluruh tahapan perancangan dan implementasi Data "
         "Warehouse, mulai dari identifikasi masalah, perancangan konseptual, "
         "logikal, hingga fisikal dengan skema Star Schema. Proses Extract, Transform, "
         "Load (ETL) dikembangkan menggunakan Python dengan memanfaatkan library "
         "Pandas, SQLAlchemy, Matplotlib, dan Seaborn untuk menghasilkan dasbor "
         "analitik yang informatif dan estetis."),

        ("Penulis menyampaikan ucapan terima kasih yang sebesar-besarnya kepada "
         "seluruh pihak yang telah memberikan dukungan dalam penyelesaian project ini, "
         "khususnya kepada dosen pengampu mata kuliah Data Warehouse yang telah "
         "memberikan bimbingan, arahan, serta data sumber yang menjadi bahan kajian "
         "dalam laporan ini."),

        ("Penulis menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena "
         "itu, kritik dan saran yang bersifat konstruktif sangat diharapkan demi "
         "penyempurnaan karya ilmiah selanjutnya. Semoga laporan ini dapat bermanfaat "
         "bagi pembaca dan memberikan kontribusi positif bagi pengembangan pengelolaan "
         "data akademik."),
    ]

    for text in pengantar_paras:
        add_body_text(doc, text, indent=True)
        doc.add_paragraph()

    # Tanda tangan
    p_ttd = doc.add_paragraph()
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ttd = p_ttd.add_run("Malang, Juni 2026")
    r_ttd.font.size = Pt(11)
    r_ttd.font.name = 'Times New Roman'

    for _ in range(4):
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(4)

    p_nama = doc.add_paragraph()
    p_nama.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_nama = p_nama.add_run("Nathanael Komang Bagus Prakarsa")
    r_nama.bold = True
    r_nama.font.size = Pt(11)
    r_nama.font.name = 'Times New Roman'

    p_nim = doc.add_paragraph()
    p_nim.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_nim = p_nim.add_run("NIM. 245091107111005")
    r_nim.font.size = Pt(11)
    r_nim.font.name = 'Times New Roman'

    add_page_break(doc)

    # =========================================================================
    # DAFTAR ISI
    # =========================================================================
    add_heading_styled(doc, "DAFTAR ISI", level=1, color_hex='1F3864')
    doc.add_paragraph()

    toc_entries = [
        ("KATA PENGANTAR", "i"),
        ("DAFTAR ISI", "ii"),
        ("DAFTAR GAMBAR", "iii"),
        ("DAFTAR TABEL", "iv"),
        ("ABSTRAK", "v"),
        ("BAB I  PENDAHULUAN", "1"),
        ("   1.1  Latar Belakang", "1"),
        ("   1.2  Rumusan Masalah", "3"),
        ("   1.3  Tujuan Penelitian", "3"),
        ("   1.4  Manfaat Penelitian", "4"),
        ("   1.5  Batasan Penelitian", "4"),
        ("BAB II  LANDASAN TEORI", "5"),
        ("   2.1  Data Warehouse", "5"),
        ("   2.2  Arsitektur Data Warehouse", "6"),
        ("   2.3  Pemodelan Dimensional", "7"),
        ("   2.4  Proses ETL (Extract, Transform, Load)", "9"),
        ("   2.5  Visualisasi Data dan Digital Dashboard", "10"),
        ("   2.6  Teknologi yang Digunakan", "11"),
        ("BAB III  METODOLOGI PENELITIAN", "12"),
        ("   3.1  Pendekatan Penelitian", "12"),
        ("   3.2  Sumber Data", "12"),
        ("   3.3  Desain Sistem Data Warehouse", "13"),
        ("   3.4  Proses ETL yang Dilakukan", "16"),
        ("   3.5  Visualisasi dan Dashboard", "18"),
        ("BAB IV  HASIL", "19"),
        ("   4.1  Hasil Implementasi Skema Database", "19"),
        ("   4.2  Hasil Proses ETL", "19"),
        ("   4.3  Dashboard Visualisasi Terpadu", "20"),
        ("   4.4  Hasil Query Analitik", "21"),
        ("BAB V  PEMBAHASAN", "26"),
        ("   5.1  Analisis Kelengkapan Data", "26"),
        ("   5.2  Analisis Performa Nilai", "27"),
        ("   5.3  Identifikasi Mahasiswa Berpotensi Risiko", "28"),
        ("   5.4  Evaluasi Implementasi DWH", "29"),
        ("DAFTAR PUSTAKA", "31"),
        ("LAMPIRAN", "33"),
        ("   Lampiran 1: DDL Script (schema.sql)", "33"),
        ("   Lampiran 2: ETL Pipeline (etl_pipeline.py)", "37"),
        ("   Lampiran 3: Query Analitik DWH", "42"),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(18)

        is_bab = entry.startswith("BAB") or entry in ["KATA PENGANTAR", "DAFTAR ISI",
                  "DAFTAR GAMBAR", "DAFTAR TABEL", "ABSTRAK", "DAFTAR PUSTAKA", "LAMPIRAN"]
        run_entry = p.add_run(entry)
        run_entry.font.size = Pt(11)
        run_entry.font.name = 'Times New Roman'
        run_entry.bold = is_bab

        # Tab filler dots
        tab = p.add_run('\t')
        tab.font.size = Pt(11)
        dots = '.' * max(1, 80 - len(entry))
        run_dots = p.add_run(dots + ' ')
        run_dots.font.size = Pt(9)
        run_dots.font.color.rgb = RGBColor(150, 150, 150)

        run_page = p.add_run(page)
        run_page.font.size = Pt(11)
        run_page.font.name = 'Times New Roman'
        run_page.bold = is_bab

    add_page_break(doc)

    # =========================================================================
    # DAFTAR GAMBAR
    # =========================================================================
    add_heading_styled(doc, "DAFTAR GAMBAR", level=1, color_hex='1F3864')
    doc.add_paragraph()

    gambar_entries = [
        ("Gambar 3.1", "Arsitektur Alur Data Warehouse Nilai Akademik", "14"),
        ("Gambar 3.2", "Star Schema - Desain Logikal Data Warehouse", "15"),
        ("Gambar 3.3", "Physical Design Star Schema (dbdiagram.io)", "16"),
        ("Gambar 4.1", "Dashboard Analitik Terpadu - Visualisasi Data Nilai Akademik", "20"),
        ("Gambar 4.2", "Chart 1: Progress Kelengkapan Data Nilai per Tipe Kelas", "21"),
        ("Gambar 4.3", "Chart 2: Distribusi Nilai Teori vs Praktikum (Violin Plot)", "22"),
        ("Gambar 4.4", "Chart 3: Rata-rata Nilai per Mata Kuliah", "23"),
        ("Gambar 4.5", "Chart 4: Distribusi Nilai per Komponen Kelas Praktikum", "24"),
        ("Gambar 4.6", "Chart 5: Peringkat Top 10 Nilai Akhir Praktikum", "25"),
    ]

    for num, desc, page in gambar_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run_n = p.add_run(f"{num}  ")
        run_n.bold = True
        run_n.font.size = Pt(11)
        run_n.font.name = 'Times New Roman'
        run_d = p.add_run(desc)
        run_d.font.size = Pt(11)
        run_d.font.name = 'Times New Roman'
        dots2 = '.' * max(1, 85 - len(num) - len(desc) - 3)
        run_dots2 = p.add_run(' ' + dots2 + ' ')
        run_dots2.font.size = Pt(9)
        run_dots2.font.color.rgb = RGBColor(150, 150, 150)
        run_p = p.add_run(page)
        run_p.font.size = Pt(11)
        run_p.font.name = 'Times New Roman'

    add_page_break(doc)

    # =========================================================================
    # DAFTAR TABEL
    # =========================================================================
    add_heading_styled(doc, "DAFTAR TABEL", level=1, color_hex='1F3864')
    doc.add_paragraph()

    tabel_entries = [
        ("Tabel 3.1", "Spesifikasi Tabel Dim_Mahasiswa", "13"),
        ("Tabel 3.2", "Spesifikasi Tabel Dim_MataKuliah", "13"),
        ("Tabel 3.3", "Spesifikasi Tabel Dim_KomponenNilai", "13"),
        ("Tabel 3.4", "Spesifikasi Tabel Dim_Periode", "14"),
        ("Tabel 3.5", "Spesifikasi Tabel Fact_Nilai", "14"),
        ("Tabel 4.1", "Hasil Query Progress Kelengkapan Data", "21"),
        ("Tabel 4.2", "Hasil Query Perbandingan Performa Teori vs Praktikum", "22"),
        ("Tabel 4.3", "Hasil Query Rata-rata Nilai per Mata Kuliah", "23"),
        ("Tabel 4.4", "Hasil Query Distribusi per Komponen Nilai", "24"),
        ("Tabel 4.5", "Identifikasi Mahasiswa dengan Nilai Rendah", "24"),
        ("Tabel 4.6", "Peringkat Nilai Akhir Estimasi Praktikum (Top 10)", "25"),
    ]

    for num, desc, page in tabel_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run_n = p.add_run(f"{num}  ")
        run_n.bold = True
        run_n.font.size = Pt(11)
        run_n.font.name = 'Times New Roman'
        run_d = p.add_run(desc)
        run_d.font.size = Pt(11)
        run_d.font.name = 'Times New Roman'
        dots3 = '.' * max(1, 85 - len(num) - len(desc) - 3)
        run_dots3 = p.add_run(' ' + dots3 + ' ')
        run_dots3.font.size = Pt(9)
        run_dots3.font.color.rgb = RGBColor(150, 150, 150)
        run_p = p.add_run(page)
        run_p.font.size = Pt(11)
        run_p.font.name = 'Times New Roman'

    add_page_break(doc)

    # =========================================================================
    # ABSTRAK
    # =========================================================================
    add_heading_styled(doc, "ABSTRAK", level=1, color_hex='1F3864')
    doc.add_paragraph()

    p_abs_title = doc.add_paragraph()
    p_abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs_title = p_abs_title.add_run(
        "IMPLEMENTASI DATA WAREHOUSE UNTUK PEMANTAUAN NILAI AKADEMIK MAHASISWA "
        "PROGRAM STUDI SAINS DATA UNIVERSITAS BRAWIJAYA"
    )
    r_abs_title.bold = True
    r_abs_title.font.size = Pt(11)
    r_abs_title.font.name = 'Times New Roman'
    p_abs_title.paragraph_format.space_after = Pt(6)

    p_abs_author = doc.add_paragraph()
    p_abs_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs_author = p_abs_author.add_run(
        "Nathanael Komang Bagus Prakarsa (245091107111005)"
    )
    r_abs_author.font.size = Pt(11)
    r_abs_author.font.name = 'Times New Roman'
    p_abs_author.paragraph_format.space_after = Pt(12)

    abstrak_text = (
        "Pengelolaan data nilai akademik mahasiswa yang tersebar dalam berbagai "
        "berkas spreadsheet (Excel) terpisah menimbulkan hambatan signifikan dalam "
        "proses pemantauan performa mahasiswa secara holistik, lintas mata kuliah, "
        "dan lintas periode akademik. Ketiadaan sistem terpadu menyulitkan dosen "
        "maupun pengelola program studi dalam mengidentifikasi tren penilaian, "
        "kemajuan koreksi, serta mendeteksi mahasiswa yang membutuhkan perhatian "
        "khusus secara dini. Penelitian ini mengimplementasikan Data Warehouse (DWH) "
        "berbasis arsitektur Kimball (Bottom-Up) dengan pemodelan dimensional Star "
        "Schema sebagai solusi integrasi data akademik. Proses Extract, Transform, "
        "Load (ETL) dikembangkan menggunakan bahasa pemrograman Python dengan "
        "memanfaatkan library Pandas untuk manipulasi data, SQLAlchemy untuk "
        "konektivitas basis data, serta SQLite sebagai RDBMS lokal. Data bersumber "
        "dari empat berkas Excel: tiga berkas presensi praktikum (MK Metode "
        "Statistika I, Metode Statistika II, dan Metode Sains Data I) serta satu "
        "berkas nilai teori (MK Metode Statistika II). Sistem yang dibangun berhasil "
        "mengkonsolidasikan 740 baris data fakta nilai dari seluruh mahasiswa kelas "
        "SA1 ke dalam satu Single Source of Truth. Hasil pipeline ETL menunjukkan "
        "bahwa 61,05% data praktikum dan 74,55% data teori telah berstatus lengkap. "
        "Visualisasi dasbor digital yang dibangun menggunakan Matplotlib dan Seaborn "
        "mengungkap rata-rata nilai praktikum sebesar 92,66 dan teori sebesar 57,25, "
        "mengindikasikan kesenjangan signifikan yang perlu mendapat perhatian dalam "
        "perencanaan pembelajaran. Implementasi Data Warehouse terbukti efektif dalam "
        "memusatkan data, mempercepat pelaporan analitik, dan menghasilkan wawasan "
        "akademik yang tidak dapat diperoleh dari spreadsheet terpisah."
    )

    p_abstrak = doc.add_paragraph()
    p_abstrak.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abstrak.paragraph_format.first_line_indent = Cm(1.25)
    p_abstrak.paragraph_format.space_after = Pt(8)
    p_abstrak.paragraph_format.line_spacing = Pt(18)
    r_abstrak = p_abstrak.add_run(abstrak_text)
    r_abstrak.font.size = Pt(11)
    r_abstrak.font.name = 'Times New Roman'

    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_kw_label = p_kw.add_run("Kata Kunci: ")
    r_kw_label.bold = True
    r_kw_label.font.size = Pt(11)
    r_kw_label.font.name = 'Times New Roman'
    r_kw_val = p_kw.add_run(
        "Data Warehouse, Star Schema, ETL, Nilai Akademik, Arsitektur Kimball, "
        "Visualisasi Data, Dashboard Analitik, Pemodelan Dimensional."
    )
    r_kw_val.font.size = Pt(11)
    r_kw_val.font.name = 'Times New Roman'
    r_kw_val.italic = True

    add_page_break(doc)

    # =========================================================================
    # BAB I: PENDAHULUAN
    # =========================================================================
    add_heading_styled(doc, "BAB I\nPENDAHULUAN", level=1, color_hex='1F3864')
    doc.add_paragraph()

    add_heading_styled(doc, "1.1 Latar Belakang", level=2, color_hex='2E5FA3')
    lb_paras = [
        ("Institusi pendidikan tinggi merupakan lingkungan yang kaya akan data. "
         "Setiap aktivitas akademik—mulai dari presensi, pengerjaan tugas, pelaksanaan "
         "kuis, ujian tengah semester, hingga ujian akhir—menghasilkan data dalam jumlah "
         "besar yang memerlukan pengelolaan sistematis. Di Program Studi Sains Data "
         "Universitas Brawijaya, data penilaian mahasiswa khususnya pada kelas Praktikum "
         "dan Teori saat ini masih dikelola secara terpisah dalam format berkas spreadsheet "
         "Microsoft Excel yang berbeda-beda untuk setiap mata kuliah dan tipe kelas."),

        ("Kondisi tersebut menimbulkan beberapa permasalahan fundamental. Pertama, "
         "fragmentasi data menyebabkan sulitnya pemantauan performa mahasiswa secara "
         "holistik dan lintas mata kuliah dalam satu periode akademik. Kedua, proses "
         "rekapitulasi nilai akhir yang harus dilakukan secara manual dari berbagai "
         "sumber berpotensi menimbulkan kesalahan dan inkonsistensi data. Ketiga, tidak "
         "terdapat visibilitas langsung terhadap kemajuan proses penilaian (berapa persen "
         "komponen nilai yang sudah dikoreksi, belum dikoreksi, atau belum dilaksanakan). "
         "Keempat, identifikasi dini terhadap mahasiswa yang mengalami kesulitan akademik "
         "menjadi terhambat karena tidak ada mekanisme analitik terintegrasi."),

        ("Data Warehouse (DWH) merupakan solusi teknologi informasi yang dirancang "
         "khusus untuk menjawab tantangan-tantangan tersebut. Berbeda dengan basis data "
         "transaksional (OLTP) yang dioptimalkan untuk operasi tulis-baca data secara "
         "cepat, Data Warehouse dioptimalkan untuk kueri analitik (OLAP) yang melibatkan "
         "agregasi, perbandingan, dan pelaporan multidimensional. Dengan mengimplementasikan "
         "Data Warehouse, seluruh data nilai akademik mahasiswa dapat dikonsolidasikan ke "
         "dalam satu repositori terpusat yang menjadi Single Source of Truth."),

        ("Project ini merespons kebutuhan tersebut dengan membangun Data Warehouse "
         "berbasis Star Schema menggunakan arsitektur Kimball (Bottom-Up). Pendekatan "
         "ini dipilih karena memungkinkan pengembangan yang cepat dan langsung berfokus "
         "pada kebutuhan analitik spesifik (Data Mart Nilai Akademik), sesuai dengan "
         "skala dan konteks kebutuhan program studi."),
    ]
    for text in lb_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "1.2 Rumusan Masalah", level=2, color_hex='2E5FA3')
    rm_intro = ("Berdasarkan latar belakang yang telah diuraikan, rumusan masalah "
                "dalam project ini adalah sebagai berikut:")
    add_body_text(doc, rm_intro, indent=True)

    rumusan = [
        ("Bagaimana merancang dan mengimplementasikan skema Data Warehouse (Star Schema) "
         "yang optimal untuk mengintegrasikan data nilai akademik mahasiswa dari berbagai "
         "sumber yang heterogen?"),
        ("Bagaimana membangun pipeline ETL yang andal dan mampu menangani anomali data "
         "(nilai kosong, format tidak standar, dan error pembagian nol) dari berkas "
         "Excel sumber?"),
        ("Bagaimana menyajikan hasil analitik Data Warehouse dalam bentuk dasbor "
         "visualisasi yang informatif untuk mendukung pengambilan keputusan akademik?"),
        ("Apa wawasan analitik yang dapat diperoleh dari Data Warehouse yang dibangun "
         "terkait performa akademik mahasiswa Program Studi Sains Data T.A. 2025/2026?"),
    ]
    for i, text in enumerate(rumusan, 1):
        add_numbered(doc, text, i)

    add_heading_styled(doc, "1.3 Tujuan Penelitian", level=2, color_hex='2E5FA3')
    tj_intro = "Tujuan yang ingin dicapai dalam project ini meliputi:"
    add_body_text(doc, tj_intro, indent=True)

    tujuan = [
        "Merancang skema dimensional Star Schema yang mencakup tabel dimensi dan fakta untuk domain nilai akademik mahasiswa.",
        "Membangun pipeline ETL berbasis Python yang mengekstrak, membersihkan, mentransformasi, dan memuat data dari berkas Excel ke dalam basis data relasional.",
        "Membuat dasbor visualisasi digital terpadu menggunakan Matplotlib dan Seaborn yang merangkum seluruh analisis dalam satu tampilan komprehensif.",
        "Mengidentifikasi pola dan wawasan akademik melalui kueri analitik pada Data Warehouse yang telah dibangun.",
    ]
    for i, text in enumerate(tujuan, 1):
        add_numbered(doc, text, i)

    add_heading_styled(doc, "1.4 Manfaat Penelitian", level=2, color_hex='2E5FA3')
    mf_intro = "Manfaat yang diharapkan dari implementasi project ini adalah:"
    add_body_text(doc, mf_intro, indent=True)

    manfaat = [
        ("Bagi Program Studi: Menyediakan sistem pemantauan nilai akademik yang terpusat, "
         "konsisten, dan dapat digunakan sebagai alat bantu pengambilan keputusan "
         "akademik berbasis data."),
        ("Bagi Dosen dan Asisten: Memberikan visibilitas langsung terhadap kemajuan "
         "penilaian, sehingga komponen nilai yang belum dikoreksi dapat teridentifikasi "
         "dengan cepat."),
        ("Bagi Mahasiswa: Memfasilitasi pemantauan perkembangan akademik secara "
         "komprehensif dan identifikasi dini terhadap kelemahan di komponen nilai tertentu."),
        ("Bagi Akademisi dan Praktisi: Menyediakan referensi implementasi Data Warehouse "
         "skala kecil dengan teknologi open-source yang dapat diadaptasi untuk konteks "
         "serupa."),
    ]
    for i, text in enumerate(manfaat, 1):
        add_numbered(doc, text, i)

    add_heading_styled(doc, "1.5 Batasan Penelitian", level=2, color_hex='2E5FA3')
    bt_intro = "Untuk menjaga fokus dan kelayakan implementasi, penelitian ini dibatasi pada:"
    add_body_text(doc, bt_intro, indent=True)

    batasan = [
        "Data yang digunakan terbatas pada empat berkas penilaian kelas SA1 untuk tiga mata kuliah pada Tahun Akademik 2025/2026.",
        "RDBMS yang digunakan adalah SQLite untuk pengujian lokal (dengan rancangan DDL yang kompatibel dengan MySQL 8.0+).",
        "Pemodelan yang digunakan terbatas pada Star Schema (tidak termasuk Snowflake Schema atau Galaxy Schema).",
        "Dashboard visualisasi dibuat dalam format gambar statis (PNG) dan tidak berbasis web interaktif.",
        "Proses ETL dijalankan secara manual melalui eksekusi skrip Python tanpa sistem penjadwalan otomatis.",
    ]
    for i, text in enumerate(batasan, 1):
        add_numbered(doc, text, i)

    add_page_break(doc)

    # =========================================================================
    # BAB II: LANDASAN TEORI
    # =========================================================================
    add_heading_styled(doc, "BAB II\nLANDASAN TEORI", level=1, color_hex='1F3864')
    doc.add_paragraph()

    add_heading_styled(doc, "2.1 Data Warehouse", level=2, color_hex='2E5FA3')
    dw_paras = [
        ("Data Warehouse (DWH) pertama kali didefinisikan secara formal oleh Bill "
         "Inmon (1990) sebagai \"koleksi data yang berorientasi subjek, terintegrasi, "
         "bervariasi terhadap waktu, dan tidak mudah berubah dalam mendukung proses "
         "pengambilan keputusan manajemen.\" Definisi ini menggarisbawahi empat "
         "karakteristik kunci yang membedakan DWH dari sistem basis data transaksional "
         "konvensional."),

        ("Pertama, berorientasi subjek (subject-oriented) berarti data DWH "
         "diorganisir berdasarkan area subjek bisnis utama seperti 'mahasiswa,' 'mata "
         "kuliah,' atau 'nilai,' bukan berdasarkan proses aplikasi tertentu. Kedua, "
         "terintegrasi (integrated) mengindikasikan bahwa data dari berbagai sumber "
         "heterogen dikonversi ke dalam format yang konsisten dan terpadu. Ketiga, "
         "bervariasi terhadap waktu (time-variant) berarti setiap rekaman data dalam "
         "DWH mengandung komponen waktu, memungkinkan analisis tren historis. Keempat, "
         "tidak mudah berubah (non-volatile) berarti data yang telah dimuat ke dalam "
         "DWH tidak diperbarui secara real-time, melainkan hanya diperbarui melalui "
         "proses ETL yang terjadwal."),

        ("Dalam konteks kecerdasan bisnis (Business Intelligence/BI), DWH berfungsi "
         "sebagai fondasi dari seluruh ekosistem pelaporan dan analitik organisasi. "
         "Dengan menyediakan Single Source of Truth, DWH memastikan bahwa seluruh "
         "pemangku kepentingan bekerja dengan data yang sama, konsisten, dan dapat "
         "dipercaya."),
    ]
    for text in dw_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "2.2 Arsitektur Data Warehouse", level=2, color_hex='2E5FA3')
    arsitektur_paras = [
        ("Terdapat dua pendekatan arsitektur utama dalam pengembangan Data Warehouse: "
         "arsitektur Inmon (Top-Down) dan arsitektur Kimball (Bottom-Up)."),

        ("Arsitektur Inmon (Top-Down) memulai pengembangan dari pembangunan Enterprise "
         "Data Warehouse (EDW) yang terpusat dan komprehensif, mencakup seluruh data "
         "organisasi. Data Mart yang spesifik untuk setiap departemen kemudian "
         "dipopulasi dari EDW. Pendekatan ini menghasilkan konsistensi data yang sangat "
         "tinggi namun memerlukan investasi waktu dan sumber daya yang signifikan di "
         "awal pengembangan."),

        ("Arsitektur Kimball (Bottom-Up), yang diadopsi dalam project ini, memulai "
         "pengembangan dengan berfokus pada pembangunan Data Mart yang spesifik untuk "
         "satu area bisnis (dalam kasus ini: Nilai Akademik Mahasiswa). Data Mart "
         "dibangun menggunakan pemodelan dimensional dan diintegrasikan secara bertahap "
         "melalui Bus Architecture. Keunggulan pendekatan ini adalah kecepatan "
         "implementasi, fokus pada kebutuhan bisnis yang spesifik, dan kemudahan "
         "penggunaan bagi pengguna akhir."),
    ]
    for text in arsitektur_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "2.3 Pemodelan Dimensional", level=2, color_hex='2E5FA3')
    pm_paras = [
        ("Pemodelan dimensional merupakan teknik desain basis data yang dikembangkan "
         "secara khusus untuk lingkungan Data Warehouse. Konsep ini diperkenalkan oleh "
         "Ralph Kimball dan dirancang untuk memaksimalkan performa kueri analitik serta "
         "kemudahan pemahaman oleh pengguna bisnis."),
    ]
    for text in pm_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "2.3.1 Tabel Fakta (Fact Table)", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Tabel fakta merupakan tabel pusat dalam model dimensional yang menyimpan "
        "pengukuran bisnis (measures atau metrics) yang dapat dikuantifikasi. Dalam "
        "konteks akademik, pengukuran tersebut adalah nilai angka mahasiswa "
        "(nilai_angka). Tabel fakta umumnya mengandung dua jenis kolom: foreign key "
        "yang merujuk ke tabel dimensi, dan angka pengukuran. Tabel fakta yang baik "
        "bersifat atomik (menyimpan data pada tingkat granularitas terendah) sehingga "
        "fleksibel untuk berbagai jenis analisis.",
        indent=True)

    add_heading_styled(doc, "2.3.2 Tabel Dimensi (Dimension Table)", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Tabel dimensi menyimpan atribut deskriptif yang memberikan konteks bagi "
        "pengukuran dalam tabel fakta. Atribut dimensi digunakan sebagai filter, "
        "grouping, dan label dalam kueri analitik. Dalam project ini, tabel dimensi "
        "meliputi Dim_Mahasiswa (siapa yang dinilai), Dim_MataKuliah (mata kuliah "
        "apa), Dim_KomponenNilai (komponen apa dan bobotnya), dan Dim_Periode (kapan "
        "penilaian berlangsung).",
        indent=True)

    add_heading_styled(doc, "2.3.3 Star Schema vs Snowflake Schema", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Star Schema adalah konfigurasi pemodelan dimensional di mana tabel fakta "
        "berada di pusat dan dikelilingi secara langsung oleh tabel dimensi yang "
        "telah didenormalisasi. Berbeda dengan Snowflake Schema yang menormalisasi "
        "dimensi ke dalam hirarki tabel tambahan, Star Schema mempertahankan semua "
        "atribut dalam satu tabel dimensi. Keunggulan Star Schema adalah: (1) "
        "performa kueri yang lebih cepat karena minimnya operasi JOIN, (2) kemudahan "
        "pemahaman oleh pengguna non-teknis, dan (3) kompatibilitas optimal dengan "
        "tool BI modern. Project ini menggunakan Star Schema karena skala data yang "
        "relatif kecil dan kebutuhan performa kueri yang tinggi.",
        indent=True)

    add_heading_styled(doc, "2.4 Proses ETL (Extract, Transform, Load)", level=2, color_hex='2E5FA3')
    etl_paras = [
        ("ETL adalah serangkaian proses yang memindahkan data dari sistem sumber ke "
         "dalam Data Warehouse. Proses ini merupakan komponen paling kritis dan "
         "seringkali paling kompleks dalam implementasi DWH."),

        ("Tahap Extract (Ekstraksi) melibatkan pembacaan data dari berbagai sistem "
         "sumber yang heterogen. Tantangan utama pada tahap ini adalah menangani "
         "perbedaan format, encoding, dan struktur antarsumber. Dalam project ini, "
         "sumber data adalah berkas Excel yang memiliki header berlebih, nilai kosong, "
         "dan error formula (#DIV/0!)."),

        ("Tahap Transform (Transformasi) adalah proses pembersihan, standarisasi, "
         "dan pengayaan data sebelum dimuat ke DWH. Transformasi mencakup: konversi "
         "tipe data, penanganan nilai null, derivasi atribut baru (seperti angkatan "
         "dari pola NIM), penerapan aturan bisnis (business rule) untuk menentukan "
         "status nilai, serta pemetaan kode sumber ke kunci surrogate DWH."),

        ("Tahap Load (Pemuatan) melibatkan penulisan data yang telah ditransformasi "
         "ke dalam tabel-tabel DWH. Strategi load yang digunakan menentukan "
         "bagaimana data baru dan data yang diperbarui ditangani. Project ini "
         "menggunakan strategi TRUNCATE-INSERT untuk tabel fakta (memastikan "
         "idempotency saat re-run) dan strategi UPSERT (ON DUPLICATE KEY UPDATE) "
         "untuk tabel dimensi."),
    ]
    for text in etl_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "2.5 Visualisasi Data dan Digital Dashboard", level=2, color_hex='2E5FA3')
    vis_paras = [
        ("Visualisasi data adalah representasi grafis dari informasi dan data. "
         "Dengan menggunakan elemen visual seperti grafik, diagram, dan peta, "
         "visualisasi memungkinkan pengguna untuk melihat dan memahami tren, "
         "outlier, dan pola dalam data yang sulit diidentifikasi dari data tabular."),

        ("Digital dashboard adalah tampilan visual yang mengkonsolidasikan dan "
         "menyajikan informasi kunci (Key Performance Indicators/KPI) dari berbagai "
         "sumber data dalam satu layar. Dalam konteks akademik, dashboard analitik "
         "memungkinkan pemangku kepentingan untuk memantau performa mahasiswa secara "
         "komprehensif sekaligus."),

        ("Library Matplotlib dan Seaborn pada Python menyediakan ekosistem "
         "visualisasi yang sangat lengkap. Matplotlib adalah library visualisasi "
         "tingkat rendah yang menyediakan kontrol penuh atas setiap elemen grafik. "
         "Seaborn, yang dibangun di atas Matplotlib, menyediakan antarmuka tingkat "
         "tinggi untuk membuat grafik statistik yang estetis dengan kode yang "
         "lebih ringkas."),
    ]
    for text in vis_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "2.6 Teknologi yang Digunakan", level=2, color_hex='2E5FA3')
    tech_paras = [
        ("Python 3.x merupakan bahasa pemrograman utama yang digunakan dalam project "
         "ini. Python dipilih karena ekosistem library data science yang sangat kaya, "
         "kemudahan sintaks, dan dukungan komunitas yang luas. Library-library kunci "
         "yang digunakan meliputi: Pandas (manipulasi dan analisis data tabular), "
         "SQLAlchemy (ORM dan konektivitas database), Matplotlib (visualisasi grafik), "
         "Seaborn (visualisasi statistik), dan python-docx (pembuatan dokumen DOCX)."),

        ("SQLite adalah RDBMS serverless yang menyimpan seluruh basis data dalam "
         "satu berkas. SQLite dipilih sebagai RDBMS lokal dalam project ini karena "
         "kemudahan setup tanpa instalasi server, performanya yang memadai untuk skala "
         "data kecil hingga menengah, dan kompatibilitas lintas platform. Rancangan "
         "DDL yang dibuat juga disiapkan untuk kompatibilitas dengan MySQL 8.0+."),

        ("MySQL adalah sistem manajemen basis data relasional open-source yang banyak "
         "digunakan dalam skala enterprise. DDL script (schema.sql) dalam project ini "
         "dirancang khusus untuk MySQL dengan fitur seperti AUTO_INCREMENT, ON "
         "DUPLICATE KEY UPDATE, dan CHECK constraints (didukung MySQL 8.0.16+)."),
    ]
    for text in tech_paras:
        add_body_text(doc, text, indent=True)

    add_page_break(doc)

    # =========================================================================
    # BAB III: METODOLOGI
    # =========================================================================
    add_heading_styled(doc, "BAB III\nMETODOLOGI PENELITIAN", level=1, color_hex='1F3864')
    doc.add_paragraph()

    add_heading_styled(doc, "3.1 Pendekatan Penelitian", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Project ini menggunakan pendekatan penelitian terapan (applied research) "
        "dengan metode implementasi sistem. Metode implementasi yang dipilih mengacu "
        "pada metodologi Kimball Lifecycle yang mencakup fase: (1) Perencanaan proyek, "
        "(2) Pendefinisian kebutuhan bisnis, (3) Desain teknikal (arsitektur, pemodelan "
        "data, desain ETL), (4) Implementasi (pengembangan ETL, basis data, dan "
        "aplikasi BI), dan (5) Deployment dan pemeliharaan.",
        indent=True)

    add_heading_styled(doc, "3.2 Sumber Data", level=2, color_hex='2E5FA3')
    sd_paras = [
        ("Data yang digunakan dalam project ini berasal dari empat berkas Microsoft "
         "Excel yang merupakan data penilaian resmi mahasiswa kelas SA1 Program Studi "
         "Sains Data, Universitas Brawijaya, Tahun Akademik 2025/2026:"),
    ]
    for text in sd_paras:
        add_body_text(doc, text, indent=True)

    sources = [
        ("Presensi Kegiatan Praktikum MK Metode Statistika I Kelas SA1.xlsx",
         "Praktikum", "Semester Ganjil 2025/2026", "Komponen: Tugas 1, Tugas 2, Presensi, Sikap, UTP (Kuis), UAP (Ujian)"),
        ("Presensi Kegiatan Praktikum MK Metode Statistika II Kelas SA1.xlsx",
         "Praktikum", "Semester Genap 2025/2026", "Komponen: Tugas 1, Tugas 2, Presensi, Sikap, UTP, UAP"),
        ("Presensi Kegiatan Praktikum MK Metode Sains Data I Kelas SA1.xlsx",
         "Praktikum", "Semester Genap 2025/2026", "Komponen: UAP (data utama yang telah terisi)"),
        ("Nilai Metode Statistika II PS Sains Data 2026.xlsx",
         "Teori", "Semester Genap 2025/2026", "Komponen: Kuis 1 (bobot 10%), Kuis 2 (10%), UTS (40%), UAS (40%)"),
    ]

    for i, (fname, tipe, periode, komponen) in enumerate(sources, 1):
        p_src = doc.add_paragraph()
        p_src.paragraph_format.space_after = Pt(4)
        p_src.paragraph_format.left_indent = Cm(1.25)
        r_num = p_src.add_run(f"{i}. ")
        r_num.bold = True
        r_num.font.size = Pt(11)
        r_num.font.name = 'Times New Roman'
        r_fname = p_src.add_run(f"{fname}\n")
        r_fname.bold = True
        r_fname.italic = True
        r_fname.font.size = Pt(10.5)
        r_fname.font.name = 'Times New Roman'
        r_detail = p_src.add_run(f"   Tipe: {tipe} | Periode: {periode}\n   {komponen}")
        r_detail.font.size = Pt(10.5)
        r_detail.font.name = 'Times New Roman'

    add_heading_styled(doc, "3.3 Desain Sistem Data Warehouse", level=2, color_hex='2E5FA3')

    add_heading_styled(doc, "3.3.1 Conceptual Design", level=3, color_hex='3A6EBE')
    cd_text = (
        "Pada tahap desain konseptual, dilakukan identifikasi dimensi dan fakta "
        "utama yang relevan dengan domain nilai akademik mahasiswa. Berdasarkan "
        "analisis kebutuhan bisnis, didefinisikan empat entitas dimensi dan satu "
        "tabel fakta sebagai berikut:"
    )
    add_body_text(doc, cd_text, indent=True)

    dim_items = [
        ("Dim_Mahasiswa", "Menyimpan informasi identitas mahasiswa (NIM, nama, kelas, angkatan)."),
        ("Dim_MataKuliah", "Menyimpan informasi mata kuliah (kode, nama, SKS)."),
        ("Dim_KomponenNilai", "Menyimpan informasi komponen penilaian (nama komponen, bobot persentase, tipe kelas)."),
        ("Dim_Periode", "Menyimpan informasi periode akademik (tahun akademik, semester)."),
        ("Fact_Nilai", "Tabel fakta yang menyimpan nilai angka mahasiswa beserta status penilaiannya, dihubungkan ke semua dimensi melalui foreign key."),
    ]
    for dim, desc in dim_items:
        p_dim = doc.add_paragraph()
        p_dim.paragraph_format.left_indent = Cm(1.25)
        p_dim.paragraph_format.space_after = Pt(3)
        r_d = p_dim.add_run(f"• {dim}: ")
        r_d.bold = True
        r_d.font.size = Pt(11)
        r_d.font.name = 'Times New Roman'
        r_desc = p_dim.add_run(desc)
        r_desc.font.size = Pt(11)
        r_desc.font.name = 'Times New Roman'

    # Gambar Alur DWH
    doc.add_paragraph()
    add_image_centered(doc, ALUR_DWH_IMG, width_inches=5.5,
                        caption="Gambar 3.1. Arsitektur Alur Data Warehouse Nilai Akademik Mahasiswa")

    add_heading_styled(doc, "3.3.2 Logical Design (Star Schema)", level=3, color_hex='3A6EBE')
    ld_text = (
        "Pada tahap desain logikal, entitas-entitas konseptual diterjemahkan ke "
        "dalam model Star Schema. Star Schema dipilih karena kemampuannya "
        "mengoptimalkan performa kueri analitik melalui denormalisasi dimensi, "
        "yang mengurangi jumlah operasi JOIN yang diperlukan. Model ini juga lebih "
        "intuitif dan mudah dipahami oleh pengguna bisnis yang melakukan kueri "
        "ad-hoc."
    )
    add_body_text(doc, ld_text, indent=True)

    add_image_centered(doc, STAR_SCHEMA_IMG, width_inches=4.5,
                        caption="Gambar 3.2. Star Schema - Desain Logikal Data Warehouse")

    add_heading_styled(doc, "3.3.3 Physical Design", level=3, color_hex='3A6EBE')
    pd_text = (
        "Pada tahap desain fisikal, model logikal diimplementasikan dalam bentuk "
        "script DDL (Data Definition Language) untuk MySQL. Setiap tabel didefinisikan "
        "dengan tipe data yang tepat, constraint (NOT NULL, UNIQUE, CHECK), dan "
        "relasi foreign key. Script ini dapat dieksekusi untuk menginisialisasi "
        "struktur basis data secara otomatis."
    )
    add_body_text(doc, pd_text, indent=True)

    # Tabel spesifikasi tabel
    def add_table_spec(doc, title, tbl_num, columns_data, header_color='1F3864'):
        """Membuat tabel spesifikasi kolom yang diformat dengan baik."""
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(f"Tabel {tbl_num}. {title}")
        r_title.bold = True
        r_title.italic = True
        r_title.font.size = Pt(10.5)
        r_title.font.name = 'Times New Roman'
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(3)

        headers = ["Nama Kolom", "Tipe Data", "Constraint", "Keterangan"]
        col_widths = [Cm(3.5), Cm(3.0), Cm(3.5), Cm(5.5)]
        table = doc.add_table(rows=1 + len(columns_data), cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = table.rows[0].cells
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            hdr[i].width = w
            hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(hdr[i], header_color)
            p_h = hdr[i].paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_h = p_h.add_run(h)
            r_h.bold = True
            r_h.font.size = Pt(9.5)
            r_h.font.color.rgb = RGBColor(255, 255, 255)
            r_h.font.name = 'Times New Roman'

        # Data rows
        for row_idx, row_data in enumerate(columns_data):
            row = table.rows[row_idx + 1]
            bg = 'F2F7FF' if row_idx % 2 == 0 else 'FFFFFF'
            for col_idx, (val, w) in enumerate(zip(row_data, col_widths)):
                cell = row.cells[col_idx]
                cell.width = w
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_bg(cell, bg)
                p_c = cell.paragraphs[0]
                p_c.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r_c = p_c.add_run(str(val))
                r_c.font.size = Pt(9.5)
                r_c.font.name = 'Courier New' if col_idx <= 1 else 'Times New Roman'
                if col_idx == 0:
                    r_c.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_table_spec(doc, "Spesifikasi Tabel Dim_Mahasiswa", "3.1", [
        ("mahasiswa_id", "INT", "PK, AUTO_INC", "Surrogate key mahasiswa"),
        ("nim", "VARCHAR(20)", "NOT NULL, UNIQUE", "Nomor Induk Mahasiswa"),
        ("nama", "VARCHAR(150)", "NOT NULL", "Nama lengkap mahasiswa"),
        ("kelas", "VARCHAR(20)", "NOT NULL", "Kode kelas (SA1, dll.)"),
        ("angkatan", "INT", "NOT NULL", "Tahun angkatan masuk"),
    ])

    add_table_spec(doc, "Spesifikasi Tabel Dim_MataKuliah", "3.2", [
        ("matkul_id", "INT", "PK, AUTO_INC", "Surrogate key mata kuliah"),
        ("kode_matkul", "VARCHAR(20)", "NOT NULL, UNIQUE", "Kode mata kuliah resmi"),
        ("nama_matkul", "VARCHAR(100)", "NOT NULL", "Nama lengkap mata kuliah"),
        ("sks", "INT", "NOT NULL", "Jumlah SKS mata kuliah"),
    ])

    add_table_spec(doc, "Spesifikasi Tabel Dim_KomponenNilai", "3.3", [
        ("komponen_id", "INT", "PK, AUTO_INC", "Surrogate key komponen"),
        ("nama_komponen", "VARCHAR(50)", "NOT NULL", "Nama komponen (Tugas 1, UTS, dst.)"),
        ("bobot_persen", "DECIMAL(5,2)", "NOT NULL", "Bobot persentase komponen"),
        ("tipe_kelas", "VARCHAR(20)", "CHECK IN (teori, praktikum)", "Tipe kelas yang bersangkutan"),
    ])

    add_table_spec(doc, "Spesifikasi Tabel Dim_Periode", "3.4", [
        ("periode_id", "INT", "PK, AUTO_INC", "Surrogate key periode"),
        ("tahun_akademik", "VARCHAR(20)", "NOT NULL", "Tahun akademik (mis. 2025/2026)"),
        ("semester", "VARCHAR(10)", "CHECK IN (Ganjil, Genap)", "Semester"),
    ])

    add_table_spec(doc, "Spesifikasi Tabel Fact_Nilai", "3.5", [
        ("nilai_id", "INT", "PK, AUTO_INC", "Surrogate key record fakta"),
        ("mahasiswa_id", "INT", "FK → Dim_Mahasiswa", "Referensi ke dimensi mahasiswa"),
        ("matkul_id", "INT", "FK → Dim_MataKuliah", "Referensi ke dimensi mata kuliah"),
        ("komponen_id", "INT", "FK → Dim_KomponenNilai", "Referensi ke dimensi komponen"),
        ("periode_id", "INT", "FK → Dim_Periode", "Referensi ke dimensi periode"),
        ("nilai_angka", "DECIMAL(5,2)", "NULL", "Nilai numerik (NULL jika belum ada)"),
        ("tipe_kelas", "VARCHAR(20)", "CHECK IN (teori, praktikum)", "Tipe kelas nilai"),
        ("status_nilai", "VARCHAR(30)", "CHECK IN (...)", "Status: lengkap/belum_dikoreksi/belum_dilaksanakan"),
    ])

    add_heading_styled(doc, "3.4 Proses ETL yang Dilakukan", level=2, color_hex='2E5FA3')

    add_heading_styled(doc, "3.4.1 Fase Extract (Ekstraksi)", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Data diekstrak dari keempat berkas Excel menggunakan fungsi pd.read_excel() "
        "dari library Pandas. Setiap berkas memiliki struktur yang berbeda: berkas "
        "praktikum membaca sheet 'Penilaian' dengan data mahasiswa dimulai dari baris "
        "ke-15 (indeks 14), sedangkan berkas teori membaca sheet-sheet terpisah "
        "('KUIS 1', 'KUIS 2', 'UTS', 'UAS') dengan data dimulai dari baris ke-3. "
        "Proses ini menggunakan parameter header=None karena baris header tidak berada "
        "di posisi standar.",
        indent=True)

    add_heading_styled(doc, "3.4.2 Fase Transform (Transformasi)", level=3, color_hex='3A6EBE')
    tf_items = [
        ("Pembersihan NIM", "Memfilter hanya baris dengan NIM yang valid (numerik) menggunakan str.isdigit(). NIM float dikonversi ke string integer untuk menghindari desimal."),
        ("Derivasi Angkatan", "Angkatan mahasiswa diderivasi dari dua digit pertama NIM dengan formula: 'angkatan = int(\"20\" + nim[:2])'."),
        ("Standarisasi Nama", "Nama mahasiswa distandardisasi ke huruf kapital semua (str.upper()) untuk konsistensi."),
        ("Penanganan Error #DIV/0!", "Sel yang mengandung string '#DIV/0!' diperlakukan sebagai nilai NULL untuk menghindari kesalahan konversi tipe data."),
        ("Rule-Based Status Nilai", "Status nilai (lengkap/belum_dikoreksi/belum_dilaksanakan) ditentukan berdasarkan aturan bisnis spesifik per mata kuliah yang mencerminkan kemajuan penilaian aktual."),
        ("Pemetaan Surrogate Key", "ID numerik dari tabel dimensi dipetakan ke dalam dictionary Python untuk lookup cepat saat pembuatan record fakta."),
    ]
    for name, desc in tf_items:
        p_tf = doc.add_paragraph()
        p_tf.paragraph_format.left_indent = Cm(1.25)
        p_tf.paragraph_format.space_after = Pt(3)
        r_n = p_tf.add_run(f"• {name}: ")
        r_n.bold = True
        r_n.font.size = Pt(11)
        r_n.font.name = 'Times New Roman'
        r_d = p_tf.add_run(desc)
        r_d.font.size = Pt(11)
        r_d.font.name = 'Times New Roman'

    add_heading_styled(doc, "3.4.3 Fase Load (Pemuatan)", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Data yang telah ditransformasi dimuat ke dalam basis data menggunakan dua "
        "strategi berbeda: (1) Strategi UPSERT (INSERT ... ON DUPLICATE KEY UPDATE) "
        "untuk tabel dimensi, memastikan data tidak terduplikasi meski script "
        "dijalankan berulang kali. (2) Strategi TRUNCATE-INSERT untuk tabel fakta "
        "Fact_Nilai, di mana data lama dihapus sepenuhnya sebelum data baru dimuat. "
        "Strategi ini menjamin idempotency—yaitu, menjalankan pipeline ETL berkali-kali "
        "akan selalu menghasilkan keadaan basis data yang sama.",
        indent=True)

    add_heading_styled(doc, "3.5 Visualisasi dan Dashboard", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Dashboard visualisasi dibangun menggunakan Matplotlib dengan GridSpec "
        "layout untuk mengatur penempatan lima panel chart dalam satu figure tunggal "
        "berukuran 22×16 inci. Desain menggunakan tema dark mode dengan palet warna "
        "neon yang dipilih secara kuratif untuk memaksimalkan keterbacaan dan estetika "
        "visual. Setiap chart dirancang untuk menjawab pertanyaan analitik yang "
        "berbeda dan saling melengkapi dalam membentuk narasi analitik yang kohesif.",
        indent=True)

    add_page_break(doc)

    # =========================================================================
    # BAB IV: HASIL
    # =========================================================================
    add_heading_styled(doc, "BAB IV\nHASIL", level=1, color_hex='1F3864')
    doc.add_paragraph()

    add_heading_styled(doc, "4.1 Hasil Implementasi Skema Database", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Skema database berhasil diinisialisasi melalui eksekusi script DDL "
        "(schema.sql) yang menciptakan empat tabel dimensi dan satu tabel fakta "
        "dengan seluruh constraint yang didefinisikan. Gambar 3.3 menunjukkan "
        "visualisasi physical design yang dibuat menggunakan dbdiagram.io, "
        "memperlihatkan relasi antar tabel beserta tipe data dan constraint "
        "masing-masing kolom.",
        indent=True)

    star_schema_ub = os.path.join(BASE_DIR, "Star Schema.png")
    add_image_centered(doc, star_schema_ub, width_inches=6.0,
                        caption="Gambar 3.3. Physical Design Star Schema - Visualisasi dbdiagram.io")

    add_heading_styled(doc, "4.2 Hasil Proses ETL", level=2, color_hex='2E5FA3')
    etl_result_paras = [
        ("Pipeline ETL berhasil dieksekusi tanpa kesalahan fatal. Proses populasi "
         "dimensi berhasil memuat 3 mata kuliah ke Dim_MataKuliah, 2 periode ke "
         "Dim_Periode, dan 10 komponen nilai ke Dim_KomponenNilai. Ekstraksi "
         "mahasiswa berhasil mengidentifikasi dan memuat seluruh mahasiswa kelas SA1 "
         "ke dalam Dim_Mahasiswa."),

        ("Fase Load menghasilkan total 740 baris data fakta di Fact_Nilai yang "
         "mencakup seluruh kombinasi mahasiswa, mata kuliah, komponen nilai, dan "
         "periode yang relevan. Dari 740 record tersebut, 516 record (69,7%) "
         "berstatus 'lengkap' dengan nilai numerik yang valid, sementara sisanya "
         "berstatus 'belum_dikoreksi' atau 'belum_dilaksanakan' mencerminkan "
         "kondisi aktual penilaian pada saat data diambil."),
    ]
    for text in etl_result_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "4.3 Dashboard Visualisasi Terpadu", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Gambar 4.1 menampilkan dashboard analitik terpadu yang merangkum lima "
        "visualisasi utama dalam satu tampilan komprehensif. Dashboard ini dirancang "
        "dengan tema dark mode premium menggunakan palet warna neon untuk "
        "memaksimalkan keterbacaan dan memberikan pengalaman visual yang modern. "
        "Bagian atas dashboard menampilkan enam KPI card yang merangkum metrik "
        "utama secara instan.",
        indent=True)

    add_image_centered(doc, DASHBOARD_IMG, width_inches=6.5,
                        caption="Gambar 4.1. Dashboard Analitik Terpadu - Visualisasi Data Nilai Akademik Mahasiswa (dark mode)")

    add_heading_styled(doc, "4.4 Hasil Query Analitik", level=2, color_hex='2E5FA3')

    # Query 1 - Progress
    add_heading_styled(doc, "4.4.1 Progress Kelengkapan Data per Tipe Kelas", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Query analitik pertama mengukur distribusi status nilai untuk setiap "
        "tipe kelas. Hasil kueri ditampilkan pada Tabel 4.1 berikut.",
        indent=True)

    def add_result_table(doc, tbl_num, title, headers, data, col_widths_cm, header_color='1F3864'):
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(f"Tabel {tbl_num}. {title}")
        r_t.bold = True
        r_t.italic = True
        r_t.font.size = Pt(10.5)
        r_t.font.name = 'Times New Roman'
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_before = Pt(6)
        p_t.paragraph_format.space_after = Pt(3)

        table = doc.add_table(rows=1 + len(data), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, (h, w) in enumerate(zip(headers, col_widths_cm)):
            hdr_cells[i].width = Cm(w)
            set_cell_bg(hdr_cells[i], header_color)
            p_h = hdr_cells[i].paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_h = p_h.add_run(h)
            r_h.bold = True
            r_h.font.size = Pt(9.5)
            r_h.font.color.rgb = RGBColor(255, 255, 255)
            r_h.font.name = 'Times New Roman'

        for ri, row_data in enumerate(data):
            row = table.rows[ri + 1]
            bg = 'EEF4FF' if ri % 2 == 0 else 'FFFFFF'
            for ci, (val, w) in enumerate(zip(row_data, col_widths_cm)):
                cell = row.cells[ci]
                cell.width = Cm(w)
                set_cell_bg(cell, bg)
                p_c = cell.paragraphs[0]
                p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_c = p_c.add_run(str(val))
                r_c.font.size = Pt(9.5)
                r_c.font.name = 'Times New Roman'

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_result_table(doc, "4.1", "Hasil Query Progress Kelengkapan Data",
        ["Tipe Kelas", "Status Nilai", "Jumlah Record", "Persentase (%)"],
        [
            ["Praktikum", "Lengkap", "315", "61.05%"],
            ["Praktikum", "Belum Dikoreksi", "201", "38.95%"],
            ["Teori", "Lengkap", "167", "74.55%"],
            ["Teori", "Belum Dilaksanakan", "56", "25.00%"],
            ["Teori", "Belum Dikoreksi", "1", "0.45%"],
        ],
        [3.5, 4.5, 3.5, 3.5]
    )

    add_image_centered(doc,
        os.path.join(BASE_DIR, "chart1_progress.png"),
        width_inches=5.5,
        caption="Gambar 4.2. Chart 1: Progress Kelengkapan Data Nilai per Tipe Kelas (Stacked Bar Chart)")

    # Query 2 - Teori vs Praktikum
    add_heading_styled(doc, "4.4.2 Perbandingan Performa Teori vs Praktikum", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Query kedua membandingkan statistik deskriptif nilai antara kelas teori "
        "dan praktikum untuk record yang berstatus 'lengkap'. Hasil kueri "
        "ditampilkan pada Tabel 4.2.",
        indent=True)

    add_result_table(doc, "4.2", "Perbandingan Statistik Nilai Teori vs Praktikum",
        ["Tipe Kelas", "Jumlah Nilai Terisi", "Rata-rata", "Nilai Min", "Nilai Maks"],
        [
            ["Praktikum", "281", "92.66", "51", "100"],
            ["Teori", "167", "57.25", "11", "100"],
        ],
        [3.0, 4.0, 3.0, 2.5, 3.0]
    )

    add_image_centered(doc,
        os.path.join(BASE_DIR, "chart2_teori_vs_praktikum.png"),
        width_inches=5.5,
        caption="Gambar 4.3. Chart 2: Distribusi Nilai Teori vs Praktikum (Violin Plot)")

    # Query 3 - per Matkul
    add_heading_styled(doc, "4.4.3 Rata-rata Nilai per Mata Kuliah", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Query ketiga mengagregasi rata-rata nilai untuk setiap kombinasi mata "
        "kuliah dan tipe kelas. Hasil kueri ditampilkan pada Tabel 4.3.",
        indent=True)

    add_result_table(doc, "4.3", "Rata-rata Nilai per Mata Kuliah dan Tipe Kelas",
        ["Mata Kuliah", "Tipe Kelas", "Jumlah Nilai Masuk", "Rata-rata Nilai"],
        [
            ["Metode Statistika II", "Praktikum", "80", "96.19"],
            ["Metode Statistika I", "Praktikum", "174", "92.38"],
            ["Metode Sains Data I", "Praktikum", "27", "84.07"],
            ["Metode Statistika II", "Teori", "167", "57.25"],
        ],
        [5.0, 3.0, 4.0, 3.5]
    )

    add_image_centered(doc,
        os.path.join(BASE_DIR, "chart3_performa_matkul.png"),
        width_inches=5.5,
        caption="Gambar 4.4. Chart 3: Rata-rata Nilai per Mata Kuliah (Grouped Bar Chart)")

    # Query 4 - Distribusi komponen
    add_heading_styled(doc, "4.4.4 Distribusi Nilai per Komponen Nilai", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Query keempat menganalisis statistik deskriptif per komponen nilai untuk "
        "setiap mata kuliah. Berikut adalah sampel hasil 12 baris pertama (Tabel 4.4).",
        indent=True)

    add_result_table(doc, "4.4", "Distribusi Nilai per Komponen (Sampel 12 Baris)",
        ["Mata Kuliah", "Komponen", "Jml Mhs", "Rata-rata", "Min", "Maks"],
        [
            ["Metode Sains Data I", "UAP", "27", "84.07", "70.00", "90"],
            ["Metode Statistika I", "Presensi", "29", "100.00", "100.00", "100"],
            ["Metode Statistika I", "Sikap", "29", "92.55", "90.00", "97"],
            ["Metode Statistika I", "Tugas 1", "29", "90.53", "55.50", "100"],
            ["Metode Statistika I", "Tugas 2", "29", "92.86", "51.00", "100"],
            ["Metode Statistika I", "UAP", "29", "92.26", "70.00", "100"],
            ["Metode Statistika I", "UTP", "29", "86.04", "63.75", "98"],
            ["Metode Statistika II", "Kuis 1", "56", "29.70", "11.00", "83"],
            ["Metode Statistika II", "Kuis 2", "55", "70.78", "18.00", "100"],
            ["Metode Statistika II", "Presensi", "28", "97.32", "62.50", "100"],
            ["Metode Statistika II", "Sikap", "28", "100.00", "100.00", "100"],
            ["Metode Statistika II", "Tugas 1", "12", "95.67", "87.00", "98"],
        ],
        [4.0, 3.0, 2.0, 2.5, 2.0, 2.0]
    )

    add_image_centered(doc,
        os.path.join(BASE_DIR, "chart4_distribusi_komponen.png"),
        width_inches=5.5,
        caption="Gambar 4.5. Chart 4: Distribusi Nilai per Komponen Kelas Praktikum (Box Plot)")

    # Query 5 - Identifikasi mahasiswa risiko
    add_heading_styled(doc, "4.4.5 Identifikasi Mahasiswa dengan Nilai Rendah", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Query kelima mengidentifikasi mahasiswa yang memiliki nilai di bawah 60 "
        "pada minimal dua komponen nilai. Hasil ini sangat penting sebagai early "
        "warning system bagi dosen untuk memberikan perhatian khusus (Tabel 4.5).",
        indent=True)

    add_result_table(doc, "4.5", "Identifikasi Mahasiswa Berpotensi Risiko Akademik (Nilai < 60, ≥ 2 Komponen)",
        ["NIM", "Nama Mahasiswa", "Mata Kuliah", "Jml Komponen Merah", "Komponen Nilai Rendah"],
        [
            ["255091101111003", "TAQIYUDDIN ELHAQ", "Metode Statistika II", "3", "Kuis 1, Kuis 2, UTS"],
            ["255091107111004", "FAIQ ADMAJA WIBOWO", "Metode Statistika II", "3", "Kuis 1, Kuis 2, UTS"],
            ["255091107111007", "AZKA NAURAH HAMASY", "Metode Statistika II", "3", "Kuis 1, Kuis 2, UTS"],
            ["255091107111014", "RAHMANIWAFDA", "Metode Statistika II", "3", "Kuis 1, Kuis 2, UTS"],
            ["255091107111016", "NOVA NABILAH PUTRI N.", "Metode Statistika II", "3", "Kuis 1, Kuis 2, UTS"],
            ["255091107111018", "AHMAD FAIZ BAIHAQY", "Metode Statistika II", "3", "Kuis 1, Kuis 2, UTS"],
            ["255091100111006", "ENCEK NAYLA AGIESTYA D.", "Metode Statistika II", "2", "Kuis 1, Kuis 2"],
            ["255091100111008", "ARYA ABIMANYU", "Metode Statistika II", "2", "Kuis 1, UTS"],
            ["255091100111013", "MUHAMMAD RAFIE HAIKAL", "Metode Statistika II", "2", "Kuis 1, Kuis 2"],
            ["255091100111017", "IDA BAGUS ARDITYA G.", "Metode Statistika I", "2", "Tugas 1, Tugas 2"],
        ],
        [3.5, 4.5, 4.0, 2.5, 4.0]
    )

    # Query 6 - Top 10
    add_heading_styled(doc, "4.4.6 Peringkat Nilai Akhir Praktikum (Top 10)", level=3, color_hex='3A6EBE')
    add_body_text(doc,
        "Query keenam menghitung nilai akhir estimasi kelas praktikum menggunakan "
        "formula rata-rata tertimbang berdasarkan bobot setiap komponen, kemudian "
        "mengurutkan 10 mahasiswa tertinggi (Tabel 4.6).",
        indent=True)

    add_result_table(doc, "4.6", "Peringkat Nilai Akhir Estimasi Praktikum (Top 10 Mahasiswa)",
        ["NIM", "Nama Mahasiswa", "Mata Kuliah", "Nilai Akhir", "Predikat"],
        [
            ["255091101111009", "AINI TUSAMMA SALSABILA", "Metode Statistika I", "98.65", "A"],
            ["255091100111012", "MELVIN TEJA", "Metode Statistika I", "97.40", "A"],
            ["255091100111005", "NATHANIEL LEKSMANA PUTRA", "Metode Statistika I", "97.20", "A"],
            ["255091100111010", "SALSABILATUS ZAHROH", "Metode Statistika I", "97.05", "A"],
            ["255091100111002", "MUHAMMAD FAIRO YANS AZFARO", "Metode Statistika I", "96.75", "A"],
            ["255091100111007", "HASRIP WIDIANTO", "Metode Statistika I", "96.55", "A"],
            ["255091100111001", "ALISTIA ESTININGTYAS", "Metode Statistika I", "96.50", "A"],
            ["255091101111002", "EONIKE GLORIA SANTOSO", "Metode Statistika I", "96.10", "A"],
            ["255091100111009", "MUHAMMAD SYAHID ALFATIH", "Metode Statistika I", "95.90", "A"],
            ["255091101111008", "MICHAEL RANDY CHRISTIAWAN", "Metode Statistika I", "95.80", "A"],
        ],
        [3.5, 5.0, 4.0, 2.5, 2.0]
    )

    add_image_centered(doc,
        os.path.join(BASE_DIR, "chart5_top_students.png"),
        width_inches=5.5,
        caption="Gambar 4.6. Chart 5: Peringkat Top 10 Nilai Akhir Praktikum (Horizontal Bar Chart)")

    add_page_break(doc)

    # =========================================================================
    # BAB V: PEMBAHASAN
    # =========================================================================
    add_heading_styled(doc, "BAB V\nPEMBAHASAN", level=1, color_hex='1F3864')
    doc.add_paragraph()

    add_heading_styled(doc, "5.1 Analisis Kelengkapan Data", level=2, color_hex='2E5FA3')
    kel_paras = [
        ("Hasil kueri progress kelengkapan data (Tabel 4.1) mengungkap kondisi "
         "penilaian yang masih dalam proses pada saat pengambilan data. Dari data "
         "praktikum, sebanyak 38,95% record masih berstatus 'belum dikoreksi.' "
         "Kondisi ini terutama disebabkan oleh komponen Tugas 2 dan UAP (Ujian Akhir "
         "Praktikum) MK Metode Statistika II yang belum selesai diinput, serta hampir "
         "seluruh komponen MK Metode Sains Data I yang baru memiliki data UAP."),

        ("Dari sisi kelas teori, 25% record berstatus 'belum dilaksanakan,' yang "
         "seluruhnya merupakan komponen UAS (Ujian Akhir Semester) MK Metode "
         "Statistika II. Hal ini mengindikasikan bahwa pada saat data diambil, UAS "
         "belum dilaksanakan, yang konsisten dengan timeline akademik Semester Genap "
         "2025/2026. Hanya 1 record (0,45%) berstatus 'belum dikoreksi' pada kelas "
         "teori, mengindikasikan hampir seluruh proses koreksi untuk komponen yang "
         "sudah dilaksanakan telah selesai."),

        ("Tingkat kelengkapan data kelas teori (74,55%) yang lebih tinggi dibanding "
         "praktikum (61,05%) mengindikasikan bahwa proses koreksi ujian teori berjalan "
         "lebih cepat relatif terhadap timeline akademik, atau bahwa komponen teori "
         "yang dijadwalkan sudah lebih banyak terlaksana dibanding komponen praktikum."),
    ]
    for text in kel_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "5.2 Analisis Performa Nilai", level=2, color_hex='2E5FA3')
    perf_paras = [
        ("Temuan paling signifikan dari analisis data adalah kesenjangan besar antara "
         "rata-rata nilai kelas praktikum (92,66) dan teori (57,25). Selisih 35,41 "
         "poin ini mengindikasikan perbedaan karakteristik penilaian yang mendasar "
         "antara kedua tipe kelas."),

        ("Pada kelas praktikum, distribusi nilai cenderung tinggi dan terkumpul pada "
         "rentang 85-100, yang tercermin dari standar deviasi yang relatif kecil. "
         "Hal ini umum pada penilaian berbasis unjuk kerja (performansi) di mana "
         "mahasiswa memiliki kesempatan untuk memperbaiki dan mengulang, serta ada "
         "komponen presensi dan sikap yang memberikan nilai penuh (100) kepada hampir "
         "semua mahasiswa."),

        ("Sebaliknya, distribusi nilai kelas teori (MK Metode Statistika II) "
         "menunjukkan variansi yang sangat besar, dengan nilai terendah 11 dan "
         "tertinggi 100. Analisis per komponen (Tabel 4.4) mengungkap bahwa Kuis 1 "
         "memiliki rata-rata sangat rendah (29,70 dari 56 mahasiswa), sementara Kuis 2 "
         "sudah jauh lebih baik (70,78). Perbedaan ekstrem ini mungkin mengindikasikan: "
         "(a) materi Kuis 1 yang lebih sulit atau belum dipelajari dengan cukup, "
         "(b) kurangnya persiapan mahasiswa pada kuis pertama, atau (c) adanya "
         "penyesuaian gaya belajar mahasiswa setelah Kuis 1."),

        ("Untuk Metode Sains Data I, data praktikum yang tersedia hanya mencakup "
         "komponen UAP dengan rata-rata 84,07. Meskipun lebih rendah dari kedua matkul "
         "statistika, nilai ini masih berada di atas ambang batas KKM yang umumnya "
         "ditetapkan pada 70-80."),
    ]
    for text in perf_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "5.3 Identifikasi Mahasiswa Berpotensi Risiko", level=2, color_hex='2E5FA3')
    risiko_paras = [
        ("Sistem DWH yang dibangun mampu mengidentifikasi mahasiswa yang memerlukan "
         "perhatian khusus secara otomatis melalui kueri analitik. Dari Tabel 4.5, "
         "teridentifikasi 10 mahasiswa yang memiliki nilai di bawah 60 pada minimal "
         "dua komponen, yang semuanya terkait dengan kelas Teori MK Metode "
         "Statistika II (kecuali satu mahasiswa dari Metode Statistika I)."),

        ("Pola yang sangat mencolok adalah enam mahasiswa yang memiliki nilai merah "
         "pada tiga komponen sekaligus (Kuis 1, Kuis 2, dan UTS). Hal ini "
         "mengindikasikan potensi risiko akademik yang serius bagi mahasiswa "
         "tersebut untuk MK Metode Statistika II. Dosen perlu memberikan perhatian "
         "dan intervensi khusus—seperti sesi konsultasi tambahan, tutorial, atau "
         "remedial—sebelum pelaksanaan UAS."),

        ("Kemampuan identifikasi dini seperti ini merupakan salah satu manfaat "
         "terbesar dari implementasi Data Warehouse. Tanpa sistem terpusat, informasi "
         "ini hanya dapat diperoleh melalui proses manual yang memakan waktu dan "
         "rentan terhadap kesalahan."),
    ]
    for text in risiko_paras:
        add_body_text(doc, text, indent=True)

    add_heading_styled(doc, "5.4 Evaluasi Implementasi Data Warehouse", level=2, color_hex='2E5FA3')

    add_heading_styled(doc, "5.4.1 Kelebihan Implementasi", level=3, color_hex='3A6EBE')
    kelebihan = [
        ("Single Source of Truth: Seluruh data nilai dari empat berkas sumber "
         "berhasil diintegrasikan ke dalam satu repositori tunggal yang konsisten "
         "dan mudah diakses."),
        ("Performa Query Optimal: Arsitektur Star Schema memungkinkan kueri "
         "analitik kompleks (melibatkan aggregasi dan join multidimensi) berjalan "
         "jauh lebih cepat dibanding memproses data Excel secara real-time."),
        ("Idempotency ETL: Pipeline ETL yang dibangun menggunakan strategi UPSERT "
         "dan TRUNCATE-INSERT memastikan konsistensi hasil meski pipeline dijalankan "
         "berulang kali."),
        ("Penanganan Anomali Data yang Andal: Sistem berhasil menangani berbagai "
         "anomali data seperti nilai null, string #DIV/0!, dan format NIM yang tidak "
         "standar tanpa menghasilkan error fatal."),
        ("Dashboard Informatif: Visualisasi terpadu memberikan gambaran komprehensif "
         "tentang kondisi penilaian dan performa akademik dalam satu tampilan."),
    ]
    for text in kelebihan:
        add_bullet(doc, text)

    add_heading_styled(doc, "5.4.2 Keterbatasan dan Saran Pengembangan", level=3, color_hex='3A6EBE')
    keterbatasan = [
        ("Tanpa Penjadwalan Otomatis: Pipeline ETL saat ini harus dijalankan secara "
         "manual. Untuk pengembangan lebih lanjut, dapat diintegrasikan dengan "
         "sistem penjadwalan seperti Apache Airflow, Luigi, atau Windows Task "
         "Scheduler."),
        ("Dashboard Statis: Visualisasi yang dihasilkan berupa gambar PNG statis, "
         "bukan dashboard web interaktif. Pengembangan menggunakan Plotly Dash, "
         "Streamlit, atau Grafana dapat memberikan pengalaman analitik yang jauh "
         "lebih kaya dan interaktif."),
        ("Cakupan Data Terbatas: Data yang digunakan hanya mencakup empat berkas "
         "dari satu semester. Sistem dapat diperluas untuk mencakup data historis "
         "dari semester-semester sebelumnya untuk analisis tren jangka panjang."),
        ("Tanpa Data Warehouse yang Sesungguhnya: Saat ini menggunakan SQLite "
         "sebagai RDBMS lokal. Untuk skala production, disarankan untuk menggunakan "
         "RDBMS yang lebih skalabel seperti MySQL/PostgreSQL atau bahkan platform "
         "cloud DWH seperti Google BigQuery atau Amazon Redshift."),
    ]
    for text in keterbatasan:
        add_bullet(doc, text)

    add_page_break(doc)

    # =========================================================================
    # DAFTAR PUSTAKA
    # =========================================================================
    add_heading_styled(doc, "DAFTAR PUSTAKA", level=1, color_hex='1F3864')
    doc.add_paragraph()

    pustaka = [
        ("Inmon, W. H. (2005). ", "Building the Data Warehouse", " (4th ed.). Wiley Publishing, Inc."),
        ("Kimball, R., & Ross, M. (2013). ", "The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling", " (3rd ed.). John Wiley & Sons, Inc."),
        ("McKinney, W. (2010). Data Structures for Statistical Computing in Python. In S. van der Walt & J. Millman (Eds.), ", "Proceedings of the 9th Python in Science Conference", " (pp. 56–61)."),
        ("Kimball, R., Reeves, L., Thornthwaite, W., & Ross, M. (2008). ", "The Data Warehouse Lifecycle Toolkit", " (2nd ed.). Wiley Publishing, Inc."),
        ("Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. ", "Computing in Science & Engineering", ", 9(3), 90–95. https://doi.org/10.1109/MCSE.2007.55"),
        ("Waskom, M. L. (2021). Seaborn: Statistical Data Visualization. ", "Journal of Open Source Software", ", 6(60), 3021. https://doi.org/10.21105/joss.03021"),
        ("SQLAlchemy Authors. (2023). ", "SQLAlchemy Documentation", ". https://docs.sqlalchemy.org/"),
        ("pandas Development Team. (2023). ", "pandas: Powerful Python Data Analysis Toolkit", ". https://pandas.pydata.org/docs/"),
        ("van der Walt, S., Colbert, S. C., & Varoquaux, G. (2011). The NumPy Array: A Structure for Efficient Numerical Computation. ", "Computing in Science & Engineering", ", 13(2), 22–30."),
        ("Oracle Corporation. (2023). ", "MySQL 8.0 Reference Manual", ". https://dev.mysql.com/doc/refman/8.0/en/"),
        ("Python Software Foundation. (2023). ", "Python 3 Documentation", ". https://docs.python.org/3/"),
        ("The SQLite Consortium. (2023). ", "SQLite Documentation", ". https://www.sqlite.org/docs.html"),
    ]

    for i, entry in enumerate(pustaka, 1):
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Cm(1.25)
        p_ref.paragraph_format.first_line_indent = Cm(-1.25)
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.line_spacing = Pt(18)

        r_num_ref = p_ref.add_run(f"{i}. ")
        r_num_ref.font.size = Pt(11)
        r_num_ref.font.name = 'Times New Roman'

        if len(entry) == 3:
            r1 = p_ref.add_run(entry[0])
            r1.font.size = Pt(11)
            r1.font.name = 'Times New Roman'
            r2 = p_ref.add_run(entry[1])
            r2.italic = True
            r2.font.size = Pt(11)
            r2.font.name = 'Times New Roman'
            r3 = p_ref.add_run(entry[2])
            r3.font.size = Pt(11)
            r3.font.name = 'Times New Roman'
        else:
            r_all = p_ref.add_run(entry[0])
            r_all.font.size = Pt(11)
            r_all.font.name = 'Times New Roman'

    add_page_break(doc)

    # =========================================================================
    # LAMPIRAN
    # =========================================================================
    add_heading_styled(doc, "LAMPIRAN", level=1, color_hex='1F3864')
    doc.add_paragraph()

    # Lampiran 1: DDL
    add_heading_styled(doc, "Lampiran 1: DDL Script - schema.sql", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Berikut adalah script DDL (Data Definition Language) lengkap yang digunakan "
        "untuk menginisialisasi struktur basis data Data Warehouse pada MySQL:",
        indent=True)

    ddl_code = """-- ==========================================
-- PHYSICAL DESIGN - DDL SCRIPT (MYSQL)
-- ==========================================
-- Tugas Project-Based Data Warehouse Ke-II
-- Topik: Data Warehouse Nilai Akademik Mahasiswa
-- ==========================================

SET FOREIGN_KEY_CHECKS = 0;

-- Hapus tabel jika sudah ada
DROP TABLE IF EXISTS Fact_Nilai;
DROP TABLE IF EXISTS Dim_Mahasiswa;
DROP TABLE IF EXISTS Dim_MataKuliah;
DROP TABLE IF EXISTS Dim_KomponenNilai;
DROP TABLE IF EXISTS Dim_Periode;

-- Tabel Dimensi: Dim_Mahasiswa
CREATE TABLE Dim_Mahasiswa (
    mahasiswa_id INT AUTO_INCREMENT PRIMARY KEY,
    nim          VARCHAR(20)  NOT NULL UNIQUE,
    nama         VARCHAR(150) NOT NULL,
    kelas        VARCHAR(20)  NOT NULL,
    angkatan     INT          NOT NULL
);

-- Tabel Dimensi: Dim_MataKuliah
CREATE TABLE Dim_MataKuliah (
    matkul_id    INT AUTO_INCREMENT PRIMARY KEY,
    kode_matkul  VARCHAR(20)  NOT NULL UNIQUE,
    nama_matkul  VARCHAR(100) NOT NULL,
    sks          INT          NOT NULL
);

-- Tabel Dimensi: Dim_KomponenNilai
CREATE TABLE Dim_KomponenNilai (
    komponen_id    INT AUTO_INCREMENT PRIMARY KEY,
    nama_komponen  VARCHAR(50)    NOT NULL,
    bobot_persen   DECIMAL(5,2)   NOT NULL,
    tipe_kelas     VARCHAR(20)    NOT NULL
        CHECK (tipe_kelas IN ('praktikum', 'teori'))
);

-- Tabel Dimensi: Dim_Periode
CREATE TABLE Dim_Periode (
    periode_id     INT AUTO_INCREMENT PRIMARY KEY,
    tahun_akademik VARCHAR(20) NOT NULL,
    semester       VARCHAR(10) NOT NULL
        CHECK (semester IN ('Ganjil', 'Genap'))
);

-- Tabel Fakta: Fact_Nilai
CREATE TABLE Fact_Nilai (
    nilai_id       INT AUTO_INCREMENT PRIMARY KEY,
    mahasiswa_id   INT          NOT NULL,
    matkul_id      INT          NOT NULL,
    komponen_id    INT          NOT NULL,
    periode_id     INT          NOT NULL,
    nilai_angka    DECIMAL(5,2) NULL,
    tipe_kelas     VARCHAR(20)  NOT NULL
        CHECK (tipe_kelas IN ('praktikum', 'teori')),
    status_nilai   VARCHAR(30)  NOT NULL
        CHECK (status_nilai IN ('lengkap','belum_dikoreksi','belum_dilaksanakan')),

    CONSTRAINT fk_mahasiswa FOREIGN KEY (mahasiswa_id)
        REFERENCES Dim_Mahasiswa(mahasiswa_id) ON DELETE CASCADE,
    CONSTRAINT fk_matkul    FOREIGN KEY (matkul_id)
        REFERENCES Dim_MataKuliah(matkul_id) ON DELETE CASCADE,
    CONSTRAINT fk_komponen  FOREIGN KEY (komponen_id)
        REFERENCES Dim_KomponenNilai(komponen_id) ON DELETE CASCADE,
    CONSTRAINT fk_periode   FOREIGN KEY (periode_id)
        REFERENCES Dim_Periode(periode_id) ON DELETE CASCADE
);

SET FOREIGN_KEY_CHECKS = 1;"""

    p_code = doc.add_paragraph()
    p_code.paragraph_format.space_after = Pt(8)
    p_code.paragraph_format.left_indent = Cm(0.5)
    r_code = p_code.add_run(ddl_code)
    r_code.font.name = 'Courier New'
    r_code.font.size = Pt(8.5)
    r_code.font.color.rgb = RGBColor(30, 30, 30)
    # Background for code block
    pPr_c = p_code._p.get_or_add_pPr()
    shd_c = OxmlElement('w:shd')
    shd_c.set(qn('w:val'), 'clear')
    shd_c.set(qn('w:color'), 'auto')
    shd_c.set(qn('w:fill'), 'F4F4F8')
    pPr_c.append(shd_c)

    doc.add_paragraph()

    # Lampiran 2: ETL
    add_heading_styled(doc, "Lampiran 2: Script ETL Pipeline - etl_pipeline.py (Cuplikan Utama)", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Berikut adalah cuplikan kode utama proses ETL yang mencakup tiga fase "
        "utama: Extract (membaca berkas Excel), Transform (pembersihan dan derivasi "
        "data), dan Load (pemuatan ke database):",
        indent=True)

    etl_snippet = """# ============================================================
# ETL PIPELINE - Fase Transform: Penentuan Status Nilai
# ============================================================
for name, (file, semester, kode_matkul) in prac_files.items():
    df = pd.read_excel(file, sheet_name='Penilaian', header=None)
    for idx, row in df.iloc[14:].iterrows():
        nim = row[5]
        if pd.notna(nim) and str(nim).strip().isdigit():
            nim_str = str(int(nim)).strip()
            for comp_name, col_idx in prac_columns_map.items():
                val = row[col_idx]
                score = None
                status = 'lengkap'
                # Tangani error #DIV/0! dan nilai kosong
                if (pd.notna(val) and str(val).strip() != ""
                        and str(val).strip() != '#DIV/0!'):
                    try:
                        score = float(val)
                    except ValueError:
                        score = None
                # Rule-based penentuan status per matkul
                if score is None:
                    if name == "MS2":
                        status = ('belum_dikoreksi'
                                  if comp_name in ["Tugas 2", "UAP"]
                                  else 'lengkap')
                    elif name == "MSD1":
                        status = ('lengkap'
                                  if comp_name == "UAP"
                                  else 'belum_dikoreksi')
                facts_buffer.append({
                    "mahasiswa_id": mahasiswa_id,
                    "matkul_id": matkul_id,
                    "komponen_id": komponen_id,
                    "periode_id": periode_id,
                    "nilai_angka": score,
                    "tipe_kelas": "praktikum",
                    "status_nilai": status
                })

# ============================================================
# Fase Load - Strategi TRUNCATE-INSERT untuk Fact Table
# ============================================================
with engine.begin() as conn:
    conn.execute(text("TRUNCATE TABLE Fact_Nilai;"))
    for fact in facts_buffer:
        conn.execute(
            text("INSERT INTO Fact_Nilai "
                 "(mahasiswa_id, matkul_id, komponen_id, "
                 "periode_id, nilai_angka, tipe_kelas, status_nilai) "
                 "VALUES (:mahasiswa_id, :matkul_id, :komponen_id, "
                 ":periode_id, :nilai_angka, :tipe_kelas, :status_nilai)"),
            fact
        )
print(f"ETL sukses! {len(facts_buffer)} baris fakta berhasil dimuat.")"""

    p_etl = doc.add_paragraph()
    p_etl.paragraph_format.space_after = Pt(8)
    p_etl.paragraph_format.left_indent = Cm(0.5)
    r_etl = p_etl.add_run(etl_snippet)
    r_etl.font.name = 'Courier New'
    r_etl.font.size = Pt(8.5)
    r_etl.font.color.rgb = RGBColor(30, 30, 30)
    pPr_e = p_etl._p.get_or_add_pPr()
    shd_e = OxmlElement('w:shd')
    shd_e.set(qn('w:val'), 'clear')
    shd_e.set(qn('w:color'), 'auto')
    shd_e.set(qn('w:fill'), 'F4F4F8')
    pPr_e.append(shd_e)

    doc.add_paragraph()

    # Lampiran 3: Query Analitik
    add_heading_styled(doc, "Lampiran 3: Query Analitik Data Warehouse", level=2, color_hex='2E5FA3')
    add_body_text(doc,
        "Berikut adalah kumpulan query SQL analitik yang dijalankan pada Data "
        "Warehouse untuk menghasilkan wawasan dari data nilai akademik:",
        indent=True)

    queries_code = """-- ====================================================
-- QUERY 1: Progress Kelengkapan Data per Tipe Kelas
-- ====================================================
SELECT
    tipe_kelas,
    status_nilai,
    COUNT(*) AS jumlah_record,
    ROUND(COUNT(*) * 100.0 /
          SUM(COUNT(*)) OVER (PARTITION BY tipe_kelas), 2) AS persentase
FROM Fact_Nilai
GROUP BY tipe_kelas, status_nilai
ORDER BY tipe_kelas, jumlah_record DESC;

-- ====================================================
-- QUERY 2: Perbandingan Performa Teori vs Praktikum
-- ====================================================
SELECT
    tipe_kelas,
    COUNT(nilai_angka)        AS jumlah_nilai_terisi,
    ROUND(AVG(nilai_angka), 2) AS rata_rata_nilai,
    MIN(nilai_angka)          AS nilai_terendah,
    MAX(nilai_angka)          AS nilai_tertinggi
FROM Fact_Nilai
WHERE status_nilai = 'lengkap'
GROUP BY tipe_kelas;

-- ====================================================
-- QUERY 3: Rata-rata Nilai per Mata Kuliah & Tipe
-- ====================================================
SELECT
    m.nama_matkul,
    f.tipe_kelas,
    COUNT(f.nilai_angka)       AS jumlah_nilai_masuk,
    ROUND(AVG(f.nilai_angka), 2) AS rata_rata_nilai
FROM Fact_Nilai f
JOIN Dim_MataKuliah m ON f.matkul_id = m.matkul_id
WHERE f.status_nilai = 'lengkap'
GROUP BY m.nama_matkul, f.tipe_kelas
ORDER BY rata_rata_nilai DESC;

-- ====================================================
-- QUERY 4: Distribusi per Komponen Nilai
-- ====================================================
SELECT
    m.nama_matkul,
    k.nama_komponen,
    COUNT(f.nilai_angka)       AS jumlah_mahasiswa,
    ROUND(AVG(f.nilai_angka), 2) AS rata_rata_nilai,
    MIN(f.nilai_angka)         AS nilai_minimal,
    MAX(f.nilai_angka)         AS nilai_maksimal
FROM Fact_Nilai f
JOIN Dim_MataKuliah m    ON f.matkul_id   = m.matkul_id
JOIN Dim_KomponenNilai k ON f.komponen_id = k.komponen_id
WHERE f.status_nilai = 'lengkap'
GROUP BY m.nama_matkul, k.nama_komponen
ORDER BY m.nama_matkul, k.nama_komponen;

-- ====================================================
-- QUERY 5: Identifikasi Mahasiswa Nilai Rendah
-- ====================================================
SELECT
    m.nim,
    m.nama,
    mk.nama_matkul,
    COUNT(f.nilai_id)                     AS jumlah_komponen_merah,
    GROUP_CONCAT(k.nama_komponen, ', ')   AS komponen_nilai_rendah
FROM Fact_Nilai f
JOIN Dim_Mahasiswa m      ON f.mahasiswa_id = m.mahasiswa_id
JOIN Dim_MataKuliah mk    ON f.matkul_id    = mk.matkul_id
JOIN Dim_KomponenNilai k  ON f.komponen_id  = k.komponen_id
WHERE f.nilai_angka < 60.00 AND f.status_nilai = 'lengkap'
GROUP BY m.nim, m.nama, mk.nama_matkul
HAVING COUNT(f.nilai_id) >= 2
ORDER BY jumlah_komponen_merah DESC;

-- ====================================================
-- QUERY 6: Peringkat Nilai Akhir Praktikum (Top 10)
-- ====================================================
SELECT
    m.nim,
    m.nama,
    mk.nama_matkul,
    ROUND(SUM(f.nilai_angka * (k.bobot_persen / 100.0)), 2)
        AS nilai_akhir_estimasi,
    CASE
        WHEN SUM(f.nilai_angka * (k.bobot_persen/100.0)) > 80 THEN 'A'
        WHEN SUM(f.nilai_angka * (k.bobot_persen/100.0)) > 75 THEN 'B+'
        WHEN SUM(f.nilai_angka * (k.bobot_persen/100.0)) > 69 THEN 'B'
        WHEN SUM(f.nilai_angka * (k.bobot_persen/100.0)) > 60 THEN 'C+'
        ELSE 'C'
    END AS predikat_estimasi
FROM Fact_Nilai f
JOIN Dim_Mahasiswa m      ON f.mahasiswa_id = m.mahasiswa_id
JOIN Dim_MataKuliah mk    ON f.matkul_id    = mk.matkul_id
JOIN Dim_KomponenNilai k  ON f.komponen_id  = k.komponen_id
WHERE f.tipe_kelas = 'praktikum' AND f.nilai_angka IS NOT NULL
GROUP BY m.nim, m.nama, mk.nama_matkul
ORDER BY nilai_akhir_estimasi DESC
LIMIT 10;"""

    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_after = Pt(8)
    p_q.paragraph_format.left_indent = Cm(0.5)
    r_q = p_q.add_run(queries_code)
    r_q.font.name = 'Courier New'
    r_q.font.size = Pt(8.5)
    r_q.font.color.rgb = RGBColor(30, 30, 30)
    pPr_q = p_q._p.get_or_add_pPr()
    shd_q = OxmlElement('w:shd')
    shd_q.set(qn('w:val'), 'clear')
    shd_q.set(qn('w:color'), 'auto')
    shd_q.set(qn('w:fill'), 'F4F4F8')
    pPr_q.append(shd_q)

    # =========================================================================
    # SIMPAN DOKUMEN
    # =========================================================================
    doc.save(OUTPUT_PATH)
    print(f"\n{'='*60}")
    print(f"LAPORAN BERHASIL DIBUAT!")
    print(f"Output: {OUTPUT_PATH}")
    print(f"{'='*60}")


# =============================================================================
# MAIN EXECUTION
# =============================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("  GENERATOR LAPORAN ILMIAH DATA WAREHOUSE")
    print("  Nathanael Komang Bagus Prakarsa - 245091107111005")
    print("=" * 60)

    print("\n[1/3] Mengunduh logo Universitas Brawijaya...")
    download_ub_logo()

    print("\n[2/3] Membangun dashboard visualisasi terpadu...")
    if os.path.exists(DB_PATH):
        build_dashboard_chart()
    else:
        print(f"  PERINGATAN: Database tidak ditemukan di {DB_PATH}")
        print("  Dashboard chart dilewati. Chart individual akan digunakan jika tersedia.")

    print("\n[3/3] Membangun dokumen laporan DOCX...")
    build_docx_report()

    print("\nSelesai! Semua proses berhasil dijalankan.")
