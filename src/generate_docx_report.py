import os
import sys
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # --- BAGIAN AWAL ---
    
    # Halaman Judul
    doc.add_heading('LAPORAN ILMIAH PROJECT DATA WAREHOUSE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Implementasi Data Warehouse Nilai Akademik Mahasiswa\n\n\n\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Disusun Oleh:').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Nathanael Komang Bagus Prakarsa\nNIM: 245091107111005').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Kata Pengantar
    doc.add_heading('Kata Pengantar', level=1)
    doc.add_paragraph('Puji syukur ke hadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya sehingga laporan project "Implementasi Data Warehouse Nilai Akademik Mahasiswa" ini dapat diselesaikan. Laporan ini disusun untuk memenuhi tugas project data warehouse. Terima kasih kepada seluruh pihak yang telah membantu dalam penyelesaian project dan penyusunan laporan ini.')
    doc.add_page_break()
    
    # Daftar Isi (Placeholder)
    doc.add_heading('Daftar Isi', level=1)
    doc.add_paragraph('1. Bagian Awal\n   - Halaman Judul\n   - Kata Pengantar\n   - Daftar Isi\n   - Abstrak\n2. Bagian Isi\n   - Pendahuluan\n   - Landasan Teori\n   - Metode\n   - Hasil\n   - Pembahasan\n3. Bagian Akhir\n   - Daftar Pustaka\n   - Lampiran')
    doc.add_page_break()
    
    # Abstrak
    doc.add_heading('Abstrak', level=1)
    doc.add_paragraph('Manajemen data nilai akademik mahasiswa seringkali terpencar dalam berbagai file spreadsheet terpisah, yang menyulitkan dosen dan institusi dalam pemantauan performa secara holistik. Project ini mengimplementasikan Data Warehouse berbasis arsitektur Kimball (Bottom-Up) dengan skema logis Star Schema. Proses Extract, Transform, Load (ETL) dibangun menggunakan Python (Pandas dan SQLAlchemy) untuk mengkonsolidasikan data dari file Excel penilaian Praktikum dan Teori ke dalam database relasional. Hasil integrasi ini kemudian divisualisasikan melalui dasbor digital untuk memantau kemajuan penilaian, distribusi komponen nilai, serta perbandingan rata-rata capaian per mata kuliah. Implementasi ini terbukti efektif dalam mempusatkan data (Single Source of Truth) dan memfasilitasi pelaporan analitik akademik secara efisien.')
    doc.add_page_break()
    
    # --- BAGIAN ISI ---
    
    # Pendahuluan
    doc.add_heading('BAB I: Pendahuluan', level=1)
    doc.add_heading('Latar Belakang', level=2)
    doc.add_paragraph('Institusi pendidikan modern menghasilkan banyak data dari aktivitas akademiknya. Permasalahan yang diangkat pada project ini adalah manajemen data nilai akademik mahasiswa yang seringkali terdistribusi pada berbagai file spreadsheet (Excel) terpisah, baik per mata kuliah, dosen, asisten, maupun tipe kelas (teori/praktikum). Data yang tersebar menyulitkan analisis performa mahasiswa secara lintas-disiplin dan monitoring kemajuan koreksi nilai. Oleh karena itu, diperlukan sebuah Data Warehouse (DWH) yang dapat bertindak sebagai Single Source of Truth.')
    doc.add_heading('Pentingnya Data Warehouse', level=2)
    doc.add_paragraph('Penerapan DWH dalam skenario ini memberikan beberapa manfaat utama: memfasilitasi pelaporan lintas mata kuliah, memberikan visibilitas langsung terhadap kemajuan penilaian dosen/asisten (seperti komponen nilai yang belum dikoreksi), dan mempercepat proses rekapitulasi nilai akhir. Tantangan utamanya adalah mengatasi format sumber data yang tidak seragam, nilai kosong, dan sel pembagian nol (misalnya #DIV/0!) melalui proses pembersihan data yang kokoh.')
    
    # Landasan Teori
    doc.add_heading('BAB II: Landasan Teori', level=1)
    doc.add_heading('Data Warehouse dan Arsitektur Kimball', level=2)
    doc.add_paragraph('Data Warehouse (DWH) adalah sistem yang digunakan untuk pelaporan dan analisis data, yang dianggap sebagai komponen inti dari kecerdasan bisnis (Business Intelligence). Arsitektur Kimball, atau pendekatan Bottom-Up, memulai pengembangan dengan berfokus pada pembangunan Data Mart (kumpulan data yang spesifik untuk area bisnis, seperti "Nilai Mahasiswa") terlebih dahulu. Data dari berbagai sumber langsung ditransformasikan ke dalam Data Mart dan diorganisir menggunakan pemodelan dimensional.')
    doc.add_heading('Skema Bintang (Star Schema)', level=2)
    doc.add_paragraph('Star Schema adalah pemodelan arsitektur logis yang mendenormalisasi struktur database menjadi satu tabel fakta pusat yang dikelilingi oleh banyak tabel dimensi terkait. Hal ini dioptimalkan untuk kueri agregat yang super cepat dibanding normalisasi basis data transaksional.')
    
    # Metode
    doc.add_heading('BAB III: Metode', level=1)
    doc.add_heading('Arsitektur yang Digunakan', level=2)
    doc.add_paragraph('Project ini menerapkan Arsitektur Bottom-Up untuk pengembangan Data Mart nilai akademik mahasiswa. Data dari berbagai file spreadsheet diekstrak dan ditransformasi langsung ke dalam skema bintang terpusat di dalam database MySQL/SQLite. Arsitektur ini sangat mendukung proses analisis performa karena query data untuk Business Intelligence hanya perlu melewati join sederhana ke satu tingkat dimensi (Online Analytical Processing/OLAP).')
    
    doc.add_heading('Desain Data Warehouse', level=2)
    doc.add_paragraph('1. Conceptual Design: Sistem mendefinisikan dimensi yang mencakup Dim_Mahasiswa (NIM, Nama), Dim_MataKuliah, Dim_KomponenNilai, dan Dim_Periode. Sedangkan untuk tabel fakta digunakan Fact_Nilai yang menampung pengukuran (nilai angka) dan foreign keys ke dimensi.\n2. Logical Design: Menggunakan pemodelan Star Schema, yang mendenormalisasi metadata agar meminimalkan operasi JOIN kompleks sehingga mempercepat rendering dashboard visual.\n3. Physical Design: Penerapan menggunakan script DDL RDBMS MySQL. Kueri pembuatan tabel utama melibatkan "CREATE TABLE" dengan deklarasi Foreign Key Constraint.')
    
    doc.add_heading('Software yang Digunakan (DWH)', level=2)
    doc.add_paragraph('Pembuatan Data Warehouse dilakukan dengan menggunakan database Relasional (SQLite / MySQL) sebagai wadah penyimpanan fisik. Pipeline pemrosesan dibangun dengan bahasa pemrograman Python melalui library SQLAlchemy untuk memetakan object database ke dalam script.')
    
    doc.add_heading('Proses Extract, Transform, Load (ETL)', level=2)
    doc.add_paragraph('Tahap ETL dijalankan menggunakan script Python dengan library Pandas. Langkah-langkah yang dilakukan:\n- Extract: Membaca sheet "Penilaian" dan "Kuis/UTS" dari file Excel mengabaikan header berlebih.\n- Transform: Membersihkan string yang tidak valid, menghapus error #DIV/0!, menghitung angkatan dari pola NIM, serta merumuskan "status_nilai" menggunakan rule-based logic.\n- Load: Mengunggah secara batch ke dalam database menggunakan fungsi ON DUPLICATE KEY untuk tabel dimensi, dan TRUNCATE-INSERT untuk load pada tabel fakta.')
    
    # Hasil
    doc.add_heading('BAB IV: Hasil', level=1)
    doc.add_heading('Implementasi Fisik dan ETL', level=2)
    doc.add_paragraph('Data sukses dikonsolidasikan dan berhasil mengisi lebih dari ratusan baris ke Fact_Nilai yang berasal dari kelas Teori Metode Statistika II serta kelas Praktikum MSD 1, MS 1, dan MS 2 tanpa adanya duplikasi.')
    
    doc.add_heading('Software Visualisasi / Digital Dashboard', level=2)
    doc.add_paragraph('Dashboard analisis dibuat menggunakan Python dengan memanfaatkan library Matplotlib dan Seaborn. Grafik-grafik disajikan dengan format "dark mode" yang premium.')
    
    doc.add_heading('Hasil Visualisasi Data', level=2)
    doc.add_paragraph('1. Progress Kelengkapan Data Nilai per Tipe Kelas: Menggunakan grafik batang bertumpuk (Stacked Bar Chart) yang merepresentasikan persentase status lengkap/belum dikoreksi.\n2. Perbandingan Performa Kelas Teori vs Praktikum: Divisualisasikan menggunakan kombinasi Violin dan Box Plot untuk mengukur variansi sentral distribusi.\n3. Rata-Rata Nilai per Mata Kuliah: Menggunakan Grouped Bar Chart untuk membedakan nilai pada kelas teori dan praktikum pada setiap mata kuliah.\n4. Distribusi Nilai per Komponen Kelas Praktikum: Disajikan melalui Box Plots untuk melihat letak rentang skor (kuartil) setiap komponen seperti UTS/UAS.\n5. Peringkat Nilai Akhir Praktikum (Top 10 Mahasiswa): Menggunakan Horizontal Bar Chart berdasar agregat tertimbang.')
    
    # Pembahasan
    doc.add_heading('BAB V: Pembahasan', level=1)
    doc.add_paragraph('Melalui penerapan arsitektur Star Schema, query pencarian untuk visualisasi performa berjalan secara sangat optimal dibandingkan memuat data langsung dari Excel secara real-time. Meskipun begitu, salah satu batasan implementasi ini adalah absennya tools penjadwalan GUI untuk ETL (seperti Pentaho/Airflow), sehingga pemicu jalannya transformasi data masih membutuhkan eksekusi command line secara manual. Sistem sudah memiliki penanganan error division by zero dan data tidak standar.')
    doc.add_page_break()
    
    # --- BAGIAN AKHIR ---
    
    # Daftar Pustaka
    doc.add_heading('Daftar Pustaka', level=1)
    doc.add_paragraph('1. Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling. Wiley.\n2. Inmon, W. H. (2005). Building the Data Warehouse. Wiley.\n3. Wes McKinney. (2010). Data Structures for Statistical Computing in Python, Proceedings of the 9th Python in Science Conference.')
    doc.add_page_break()
    
    # Lampiran
    doc.add_heading('Lampiran', level=1)
    doc.add_paragraph('Lampiran 1: Kode DDL Pembuatan Tabel (schema.sql)\nLampiran 2: Cuplikan Script ETL Pipeline (etl_pipeline.py)\nLampiran 3: Screenshot Dashboard Visualisasi yang di-generate.')
    
    # Simpan File
    output_path = r'd:\Nael\Mata Kuliah Sains Data\.vscode\Kuliah\Semester 4\DWH Dosen\Laporan_Ilmiah_Data_Warehouse.docx'
    doc.save(output_path)
    print(f"File laporan DOCX berhasil dibuat di: {output_path}")

if __name__ == '__main__':
    create_report()
